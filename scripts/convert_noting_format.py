import os
import docx
from docx.shared import Pt, Inches

def convert_document_with_formatting():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    src_path = os.path.join(base_dir, 'masters', 'Common_Noting.docx')
    dst_path = os.path.join(base_dir, 'masters', 'TBRL_Noting_Master.docx')
    
    if not os.path.exists(src_path):
        print(f"Error: {src_path} does not exist!")
        return
        
    doc = docx.Document(src_path)
    
    # ── Paragraph 0: Reference number & Date ──
    # Runs: [('संदर्भ सं. ', True), ('टीबीआरएल/मा.सं.वि.वि./    ', None), ('                                    दिनांक:   /    /2025  ', True)]
    p0 = doc.paragraphs[0]
    if len(p0.runs) >= 3:
        p0.runs[1].text = " {{LAB_NAME}}/{{REF_NO}} "
        p0.runs[2].text = p0.runs[2].text.replace("2025", "{{CURRENT_YEAR}}")
        
    # ── Paragraph 2: Subject Line ──
    # Keep run 0 (विषय) and run 1 (/Sub :), set run 2 to placeholder and clear the rest
    p2 = doc.paragraphs[2]
    if len(p2.runs) >= 3:
        p2.runs[2].text = " {{SUBJECT_HINDI}} / {{SUBJECT_ENGLISH}}"
        for r in p2.runs[3:]:
            r.text = ""
        
    # ── Paragraph 4: Reference ──
    # Keep run 0 ( संदर्भ : ) and replace the rest with placeholder
    p4 = doc.paragraphs[4]
    if len(p4.runs) >= 2:
        p4.runs[1].text = " {{REFERENCE_TEXT}} ..(1)"
        for r in p4.runs[2:]:
            r.text = ""

    # ── Paragraph 7: Intro Text ──
    # Set Run 0 to placeholder text and clear the rest to keep font settings
    p7 = doc.paragraphs[7]
    if len(p7.runs) >= 1:
        p7.runs[0].text = "2. उपरोक्त संलग्न संदर्भ सं. (1), के अनुसार {{COURSE_DATES}} के दौरान {{ORG_INSTITUTE}} द्वारा, {{LAB_NAME}} में {{COURSE_TITLE}} विषय पर {{COURSE_TYPE}} आयोजित किया जा रहा है। इस कोर्स में भाग लेने के लिए प्राप्त {{OFFICIAL_WORD}} {{KA_KE}} {{NAMANKAN_WORD}} का विवरण निम्नलिखित है।"
        for r in p7.runs[1:]:
            r.text = ""

    # ── Paragraph 10: Request Text ──
    # Set Run 0 to placeholder text and clear the rest
    p10 = doc.paragraphs[10]
    if len(p10.runs) >= 1:
        p10.runs[0].text = "4.आपसे निवेदन है कि प्राप्त {{NAMANKAN_WORD}} से या किसी अन्य उपयुक्त ऑफिशियल को उपरोक्त विषय में भाग लेने के लिए नामांकित करें।"
        for r in p10.runs[1:]:
            r.text = ""

    # ── Signatories Paragraphs ──
    # We target the exact paragraph indices for signatories to keep formatting
    p17 = doc.paragraphs[17]
    for run in p17.runs:
        if "राकेश कुमार" in run.text:
            run.text = run.text.replace("राकेश कुमार", "{{SIGNATORY_1_NAME}}")
            
    p18 = doc.paragraphs[18]
    for run in p18.runs:
        if "वैज्ञानिक ‘एफ’(मा.सं.वि.वि)" in run.text:
            run.text = run.text.replace("वैज्ञानिक ‘एफ’(मा.सं.वि.वि)", "{{SIGNATORY_1_DESIG}}")
            
    p21 = doc.paragraphs[21]
    for run in p21.runs:
        if "डॉ. अनिल  खुराना, वैज्ञानिक 'जी'" in run.text:
            run.text = run.text.replace("डॉ. अनिल  खुराना, वैज्ञानिक 'जी'", "{{SIGNATORY_2_NAME}}")
        elif "डॉ. अनिल" in run.text:
            run.text = run.text.replace(run.text, "  {{SIGNATORY_2_NAME}}")
            
    p22 = doc.paragraphs[22]
    for run in p22.runs:
        if "ग्रुप निदेशक(मा.सं.वि.वि.)" in run.text:
            run.text = run.text.replace("ग्रुप निदेशक(मा.सं.वि.वि.)", "{{SIGNATORY_2_DESIG}}")
        elif "ग्रुप निदेशक" in run.text:
            run.text = run.text.replace(run.text, "       {{SIGNATORY_2_DESIG}}")

    # ── Table replacement ──
    # Locate first table, remove it, insert {{COMPLEX_DYNAMIC_TABLE}} placeholder paragraph
    if len(doc.tables) > 0:
        table_to_replace = doc.tables[0]
        tbl_element = table_to_replace._tbl
        parent = tbl_element.getparent()
        
        new_para = doc.add_paragraph()
        new_para.text = "{{COMPLEX_DYNAMIC_TABLE}}"
        
        tbl_element.addprevious(new_para._element)
        parent.remove(tbl_element)
        print("✔ Table successfully replaced with {{COMPLEX_DYNAMIC_TABLE}}.")

        if os.path.exists(dst_path):
        import shutil
        from datetime import datetime
        backup_path = dst_path + f".bak_{datetime.now():%Y%m%d%H%M%S}"
        shutil.copy2(dst_path, backup_path)
        print(f"✔ Backup created at: {backup_path}")
    doc.save(dst_path)
    print(f"✔ Formatting preserved template master generated at: {dst_path}")


if __name__ == '__main__':
    convert_document_with_formatting()
