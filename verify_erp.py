import sys
import os
import docx

# Import generator
from logic.noting_generator import generate_mayurpankh_erp_fax

mock_data = {
    "ref_no": "टीबीआरएल/एचआरडीडी/2026/089",
    "fax_no": "020-27044202,27044010",
    "ref_date": "17/06/2026",
    "to_text": "निदेशक महोदय (The Director)\nअनुसंधानतथाविकास स्थापन (इंजीनियर्स) (Research & Development Establishment (Engineers))\nकलस, आलंदी रोड, दिघी पो.ऑ. (Kalas, Alandi Road,Dighi P.O.)\nपुणे-411015 (Pune-411015)",
    "subject_hindi": "\"DRDO MAYURPANKH ERP\" सी.ई.पी  कोर्स के लिए नामांकन ।",
    "subject_english": "Nomination for CEP Course on \"DRDO MAYURPANKH ERP\"",
    "ref_text": "RDE/90001/CEP/IMSG      Dated  03/07/2024",
    "attn_text": "सुश्री . गुरप्रीत कौर, वैज्ञानिक, 'एफ़'/ Ms. Gurpreet Kaur, Scientist ‘F’)\nग्रुप निदेशक, आई.एम.एस.जी/एच.आर.डी. / GD, IMSG/ HRD)",
    "body_hindi": "आर&डी.ई.(ई), पुणे , द्वारा 20 से 23 अगस्त 2024 को पुणे मे \"DRDO MAYURPANKH ERP\" विषय पर  होने वाले  सी.ई.पी. कोर्स में भाग लेने  हेतु निदेशक महोदय के द्वारा निम्नलिखित प्रतिभागी के नाम अनुमोदित किया गया हैं:-",
    "body_english": "The nomination of the following official have been approved by the competent authority to attend the CEP Course on “DRDO MAYURPANKH ERP” scheduled on 20-23 August 2024 at R&DE (E), Pune.",
    "confirm_text": "You are requested to confirm the above nomination through return DRONA mail hrd@tbrl.chddom/ Fax (0172-2308105) and arrange necessary arrangements for the confirmed participant.",
    "columns": ["PIS", "Name /Rank", "DOB", "Contact No./ Email ID", "Qualification", "Key Functional Area"],
    "rows": [
        {
            "PIS": "2014DA1009",
            "Name /Rank": "Sh. Neeraj Kumar, Accountant",
            "DOB": "01.09.1986",
            "Contact No./ Email ID": "9911420419",
            "Qualification": "B.A.",
            "Key Functional Area": "Finance"
        }
    ],
    "sig_name": "राकेश कुमार / Rakesh Kumar",
    "sig_desig": "वैज्ञानिक, 'ई' /Scientist ‘E’",
    "for_director": "कृते निदेशक/ For Director"
}

print("Testing generate_mayurpankh_erp_fax...")
try:
    filepath, filename = generate_mayurpankh_erp_fax(mock_data)
    print(f"✔ File generated at: {filepath}")
    
    # Verify file contents
    doc = docx.Document(filepath)
    has_unreplaced_placeholders = False
    
    print("\n--- Document Paragraphs ---")
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"P{idx}: {para.text}")
            if "{{" in para.text or "}}" in para.text:
                has_unreplaced_placeholders = True
                print(f"  ❌ UNREPLACED PLACEHOLDER: {para.text}")
                
    print("\n--- Document Tables ---")
    for idx, table in enumerate(doc.tables):
        print(f"Table {idx}:")
        for r_idx, row in enumerate(table.rows):
            cells_txt = [c.text.strip().replace("\n", " ") for c in row.cells]
            print(f"  Row {r_idx}: {cells_txt}")
            
    if has_unreplaced_placeholders:
        print("\n❌ FAILED: Found unreplaced placeholders!")
        sys.exit(1)
    else:
        print("\n🎉 SUCCESS: All placeholders replaced successfully!")
        sys.exit(0)
        
except Exception as e:
    import traceback
    traceback.print_exc()
    sys.exit(1)
