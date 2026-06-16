import sys
import os
import docx

# Import generator
from logic.noting_generator import generate_date_amendment_fax

mock_data = {
    "ref_no": "टीबीआरएल/एचआरडीडी/2026/045",
    "fax_no": "011-23014903",
    "ref_date": "17/06/2026",
    "to_text": "निदेशक महोदय/ (The Director)\nरक्षा अनुसंधान तथा विकास संगठन/ (Defence R&D Organisation)\nमानव संसाधन विकास निदेशालय 'बी ब्लाक’/ (Dte. Of Human Resource Dev)\n'बी ब्लाक’ डी.आर.डी.ओ, भवन/ ('B' Wing, DRDO Bhawan)\nराजाजी मार्ग, नई दिल्ली -110011/ (Rajaji Marg, New Delhi- 110011)",
    "subject_hindi": "सी.ई.पी. कोर्स के आयोजन तिथि में संशोधन",
    "subject_english": "Amendment in Date of CEP Courses",
    "ref_text": "सी.ई.पी. हैंड्बुक 2024-25(CEP Handbook 2024-25),",
    "ref_to_hindi": "सी.ई.पी. हैंड्बुक 2024-25",
    "ref_to_eng": "CEP Handbook 2024-25",
    "num_courses_hindi": "तीन",
    "num_courses_eng": "three",
    "courses_plural_eng": "Courses",
    "amended_courses_hindi": "एक",
    "amended_courses_eng": "one",
    "for_columns": ["S. No.", "Course Name", "Lab", "Month", "Duration", "Eligibility Criteria"],
    "for_rows": [
        {
            "Course Name": "High Explosive Technology",
            "Lab": "TBRL,",
            "Month": "May 2024",
            "Duration": "20 से 24 मई 2024",
            "Eligibility Criteria": "DRDS&DRTC"
        }
    ],
    "read_columns": ["S. No.", "Course Name", "Lab", "Month", "Duration", "Eligibility Criteria"],
    "read_rows": [
        {
            "Course Name": "High Explosive Technology",
            "Lab": "TBRL,",
            "Month": "June 2024",
            "Duration": "10 से 14 जून 2024",
            "Eligibility Criteria": "DRDS&DRTC"
        }
    ],
    "sig_name": "राकेश कुमार /Rakesh Kumar",
    "sig_desig": "वैज्ञानिक, 'ई' /Scientist ‘E’",
    "for_director": "कृते निदेशक/For Director"
}

print("Testing generate_date_amendment_fax...")
try:
    filepath, filename = generate_date_amendment_fax(mock_data)
    print(f"✔ File generated at: {filepath}")
    
    # Verify file contents
    doc = docx.Document(filepath)
    has_unreplaced_placeholders = False
    
    print("\n--- Document Paragraphs ---")
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"P{idx}: {para.text}")
            if "{{" in para.text or "}}" in para.text:
                has_unreplaced_placeholders = True
                print(f"  ❌ UNREPLACED PLACEHOLDER: {para.text}")
                
    print("\n--- Document Tables ---")
    for idx, table in enumerate(doc.tables):
        print(f"Table {idx}:")
        for r_idx, row in enumerate(table.rows):
            cells_txt = [c.text.strip().replace("\n", " ") for c in row.cells]
            print(f"  Row {r_idx}: {cells_txt}")
            
    if has_unreplaced_placeholders:
        print("\n❌ FAILED: Found unreplaced placeholders!")
        sys.exit(1)
    else:
        print("\n🎉 SUCCESS: All placeholders replaced successfully!")
        sys.exit(0)
        
except Exception as e:
    import traceback
    traceback.print_exc()
    sys.exit(1)
