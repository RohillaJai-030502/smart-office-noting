import os
import docx

def convert_fee_document():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    src_path = os.path.join(base_dir, 'masters', 'For_FeeRelated_Noting.docx')
    dst_path = os.path.join(base_dir, 'masters', 'Fee_Noting_Master.docx')
    
    if not os.path.exists(src_path):
        print(f"Error: {src_path} does not exist!")
        return
        
    doc = docx.Document(src_path)
    
    # ── Paragraph 0: Reference number & Date ──
    p0 = doc.paragraphs[0]
    if len(p0.runs) >= 3:
        p0.runs[1].text = " {{LAB_NAME}}/{{REF_NO}} "
        p0.runs[2].text = "                                                                दिनांक: {{REF_DATE}}"
        
    # ── Paragraph 1: Subject Line ──
    p1 = doc.paragraphs[1]
    if len(p1.runs) >= 5:
        p1.runs[1].text = " {{SUBJECT_HINDI}}"
        p1.runs[2].text = ""
        p1.runs[3].text = ""
        p1.runs[4].text = " ({{SUBJECT_ENGLISH}})"

    # ── Paragraph 3: Reference ──
    p3 = doc.paragraphs[3]
    if len(p3.runs) >= 4:
        p3.runs[1].text = ": {{REFERENCE_TEXT}}                             "
        p3.runs[3].text = ": {{REF_MAIL_DATE}}……...(1)"

    # ── Paragraph 5: Intro Text ──
    p5 = doc.paragraphs[5]
    if len(p5.runs) >= 1:
        p5.runs[0].text = "2. उपरोक्त संलग्न संदर्भ सं. (1), के अनुसार {{COURSE_DATES}} के दौरान {{ORG_INSTITUTE}} द्वारा, {{LAB_NAME}} में {{COURSE_TITLE}} विषय पर {{COURSE_TYPE}} आयोजित किया जा रहा है। इस कोर्स में भाग लेने के लिए प्राप्त {{OFFICIAL_WORD}} {{KA_KE}} {{NAMANKAN_WORD}} का विवरण निम्नलिखित है।"
        for r in p5.runs[1:]:
            r.text = ""

    # ── Signatories Paragraphs ──
    p12 = doc.paragraphs[12]
    for run in p12.runs:
        if "राकेश कुमार" in run.text:
            run.text = run.text.replace("राकेश कुमार", "{{SIGNATORY_1_NAME}}")
            
    p13 = doc.paragraphs[13]
    for run in p13.runs:
        if "वैज्ञानिक ‘एफ’(मा.सं.वि.वि)" in run.text:
            run.text = run.text.replace("वैज्ञानिक ‘एफ’(मा.सं.वि.वि)", "{{SIGNATORY_1_DESIG}}")
            
    p16 = doc.paragraphs[16]
    for run in p16.runs:
        if "डॉ. अनिल  खुराना, वैज्ञानिक 'जी'" in run.text:
            run.text = run.text.replace("डॉ. अनिल  खुराना, वैज्ञानिक 'जी'", "{{SIGNATORY_2_NAME}}")
        elif "डॉ. अनिल" in run.text:
            run.text = run.text.replace(run.text, "  {{SIGNATORY_2_NAME}}")
            
    p17 = doc.paragraphs[17]
    for run in p17.runs:
        if "ग्रुप निदेशक(मा.सं.वि.वि.)" in run.text:
            run.text = run.text.replace("ग्रुप निदेशक(मा.सं.वि.वि.)", "{{SIGNATORY_2_DESIG}}")
        elif "ग्रुप निदेशक" in run.text:
            run.text = run.text.replace(run.text, "       {{SIGNATORY_2_DESIG}}")

    # ── Table replacement ──
    if len(doc.tables) > 0:
        table_to_replace = doc.tables[0]
        tbl_element = table_to_replace._tbl
        parent = tbl_element.getparent()
        
        new_para = doc.add_paragraph()
        new_para.text = "{{COMPLEX_DYNAMIC_TABLE}}"
        
        tbl_element.addprevious(new_para._element)
        parent.remove(tbl_element)
        print("✔ Table successfully replaced with {{COMPLEX_DYNAMIC_TABLE}}.")

    doc.save(dst_path)
    print(f"✔ Formatting preserved template master generated at: {dst_path}")

if __name__ == '__main__':
    convert_fee_document()
