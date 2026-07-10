import re
import uuid
from datetime import datetime
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
from docx.oxml.table import CT_Tbl

def get_unique_timestamp():
    """Generates a timestamp suffixed with a unique 6-character UUID segment to avoid collisions."""
    return datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:6]


def get_safe(d, key, default=""):
    """Returns the default if the key is missing OR if the value is None."""
    if d is None:
        return default
    val = d.get(key, default)
    return default if val is None else val

def parse_date_safe(value, fmt="%Y-%m-%d"):
    """Parses a date string safely, returning a datetime object or raising ValueError."""
    if not value:
        raise ValueError("Missing or empty date value.")
    try:
        return datetime.strptime(str(value).strip(), fmt)
    except (TypeError, ValueError) as e:
        raise ValueError(f"Invalid date format: '{value}'. Expected '{fmt}'.") from e

def replace_placeholders_in_paragraph(paragraph, placeholders):
    """Handles placeholders split across multiple runs by Microsoft Word."""
    if not paragraph.runs:
        return
    
    full_text = "".join(run.text for run in paragraph.runs)
    original = full_text
    
    for key, value in placeholders.items():
        if key in full_text:
            full_text = full_text.replace(key, str(value))
            
    if full_text != original:
        # Put the fully-replaced text into the first run
        paragraph.runs[0].text = full_text
        # Empty out all subsequent runs in the paragraph to avoid duplication
        for run in paragraph.runs[1:]:
            run.text = ""

def has_unreplaced_placeholders(doc):
    """Scans doc paragraphs and tables for any unreplaced {{...}} placeholders."""
    pattern = re.compile(r"\{\{.+?\}\}")
    
    for para in doc.paragraphs:
        if pattern.search(para.text):
            return True
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if pattern.search(para.text):
                        return True
                        
    return False


def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    """Sets compact cell padding (in dxas) for tighter table density."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table, color="CCCCCC", sz="4"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def apply_table_formatting(table):
    """Enforces cantSplit on all rows, tblHeader on the first row, and explicit column widths/padding."""
    if len(table.rows) > 0:
        trPr = table.rows[0]._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

    # Dynamic column widths calculation to prevent overflow (A4 printable area = 6.75 inches)
    num_cols = len(table.columns)
    if num_cols > 0:
        total_width = Inches(6.75)
        widths = [total_width / num_cols] * num_cols
        
        # Read headers
        headers = []
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip().lower())
                
        # Assign specific widths based on keywords
        assigned_total = 0
        assigned_indices = set()
        
        for idx, header in enumerate(headers):
            if any(kw in header for kw in ["s.n.", "sl.", "क्रम"]):
                widths[idx] = Inches(0.4)
                assigned_total += Inches(0.4)
                assigned_indices.add(idx)
            elif any(kw in header for kw in ["pis", "roll", "mobile", "contact", "phone", "फोन", "रोल"]):
                widths[idx] = Inches(0.9)
                assigned_total += Inches(0.9)
                assigned_indices.add(idx)
            elif any(kw in header for kw in ["name", "design", "college", "institute", "संगठन", "नाम", "पद"]):
                widths[idx] = Inches(1.8)
                assigned_total += Inches(1.8)
                assigned_indices.add(idx)

        # Distribute remaining width
        unassigned_count = num_cols - len(assigned_indices)
        if unassigned_count > 0:
            remaining_width = max(Inches(1.0), total_width - assigned_total)
            share = remaining_width / unassigned_count
            for idx in range(num_cols):
                if idx not in assigned_indices:
                    widths[idx] = share
                    
        # Apply cell widths
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(widths):
                    cell.width = widths[idx]

    # Set tight global cell margins (padding) on the table
    tblPr = table._tbl.tblPr
    existing_mar = tblPr.find(qn('w:tblCellMar'))
    if existing_mar is not None:
        tblPr.remove(existing_mar)
    tblPr.append(parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="80" w:type="dxa"/>'
        f'  <w:bottom w:w="80" w:type="dxa"/>'
        f'  <w:left w:w="120" w:type="dxa"/>'
        f'  <w:right w:w="120" w:type="dxa"/>'
        f'</w:tblCellMar>'
    ))


def build_dynamic_table_in_doc(doc, para, columns, rows):
    """Creates a print-ready formatted dynamic table in place of the target paragraph."""
    parent = para._element.getparent()
    tbl = CT_Tbl.new_tbl(len(rows) + 1, len(columns), Inches(6.5))
    para_index = parent.index(para._element)
    parent.insert(para_index + 1, tbl)
    
    table = Table(tbl, para._parent)
    
    set_table_borders(table, color="A0AEC0", sz="4")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 1. Fill headers
    hdr_cells = table.rows[0].cells
    for col_idx, col_name in enumerate(columns):
        cell = hdr_cells[col_idx]
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        run = p.add_run(col_name)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        run.font.size = Pt(9.5)
        
    # 2. Fill data rows
    sn_aliases = ["s.n.", "s.no.", "sl.no.", "cr.sn.", "क्र.सं.", "क्र.स", "क्र.सं", "क्र.स."]
    for row_idx, row_data in enumerate(rows, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, col_name in enumerate(columns):
            cell = row_cells[col_idx]
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            is_sn = any(alias in col_name.lower() for alias in sn_aliases)
            if is_sn:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(row_idx))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9.5)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                val = row_data.get(col_name, "")
                run = p.add_run(str(val))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9.5)
                
    # 3. Format widths
    apply_table_formatting(table)
    
    # 4. Remove original paragraph containing placeholder
    parent.remove(para._element)
