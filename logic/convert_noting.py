import os
import docx

def convert_document():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    src_path = os.path.join(base_dir, 'masters', 'Common_Noting.docx')
    dst_path = os.path.join(base_dir, 'masters', 'TBRL_Noting_Master.docx')
    
    if not os.path.exists(src_path):
        print(f"Error: {src_path} does not exist!")
        return
        
    doc = docx.Document(src_path)
    
    # 1. Paragraph replacement mapping
    # Note: Using substrings to make sure we match runs safely
    replacements = [
        # Reference number line
        ("संदर्भ सं. टीबीआरएल/मा.सं.वि.वि./", "संदर्भ सं. {{LAB_NAME}}/{{REF_NO}}"),
        ("दिनांक:   /    /2025", "दिनांक:   /    /{{CURRENT_YEAR}}"),
        
        # Subject line (handled via startswith below)
        # Reference line
        ("Received mail from “Director PM” <Dir_pm@dgecs.bandom >  Dated: 26th  Nov 2025 ..(1)", "{{REFERENCE_TEXT}} ..(1)"),
        
        # Body paragraph
        ("उपरोक्त संलग्न संदर्भ सं. (1), के अनुसार 20-22 जनवरी के दौरानआई.आई.एस.सी.,बेंगलोरु (IISc,Bangalore,India) द्वारा, TBRL में \"इलेक्ट्रॉनिक वारफेयर ई.डब्लू.सी.आई. 2026 \" विषय पर कॉन्फेरेंस / कार्यशाला /सी.ई.पी. कोर्स/संगोष्ठी/एम.डी.पी. कोर्स/सिम्पोजियम/ आयोजित किया  जा रहा  है। इस कोर्स में भाग लेने के लिएप्राप्त नामांकन /नामांकनों/ का विवरण निम्नलिखित है।",
         "उपरोक्त संलग्न संदर्भ सं. (1), के अनुसार {{COURSE_DATES}} के दौरान {{ORG_INSTITUTE}} द्वारा, {{LAB_NAME}} में {{COURSE_TITLE}} विषय पर {{COURSE_TYPE}} आयोजित किया जा रहा है। इस कोर्स में भाग लेने के लिए प्राप्त {{OFFICIAL_WORD}} {{KA_KE}} {{NAMANKAN_WORD}} का विवरण निम्नलिखित है।"),
         
        # Body paragraph 4 & 5
        ("प्राप्त नामांकन से या किसी अन्य उपयुक्तऑफिसियल को उपरोक्त विषय में भाग लेने  के लिए नामांकित करें।", "प्राप्त {{NAMANKAN_WORD}} से या किसी अन्य उपयुक्त ऑफिशियल को उपरोक्त विषय में भाग लेने के लिए नामांकित करें।"),
        
        # Signatories
        ("राकेश कुमार", "{{SIGNATORY_1_NAME}}"),
        ("वैज्ञानिक ‘एफ’(मा.सं.वि.वि)", "{{SIGNATORY_1_DESIG}}"),
        ("डॉ. अनिल  खुराना, वैज्ञानिक 'जी'", "{{SIGNATORY_2_NAME}}"),
        ("ग्रुप निदेशक(मा.सं.वि.वि.)", "{{SIGNATORY_2_DESIG}}")
    ]
    
    # Apply replacements on paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("विषय/Sub :") or text.startswith("विषय/Sub:") or text.startswith("विषय /Sub"):
            text = "विषय/Sub : {{SUBJECT_HINDI}} / {{SUBJECT_ENGLISH}}"
        else:
            for old_txt, new_txt in replacements:
                if old_txt in text:
                    text = text.replace(old_txt, new_txt)
        para.text = text
        
    # 2. Table replacement: Find the table, delete it, and replace with {{COMPLEX_DYNAMIC_TABLE}} placeholder paragraph
    # We locate the table element in the document body
    body = doc.element.body
    table_index = None
    
    # We find the table in doc.tables
    for idx, table in enumerate(doc.tables):
        # We will replace the first table we find
        table_index = idx
        break
        
    if table_index is not None:
        table_to_replace = doc.tables[table_index]
        # Get its XML element
        tbl_element = table_to_replace._tbl
        # Find parent
        parent = tbl_element.getparent()
        
        # Create a new paragraph element for {{COMPLEX_DYNAMIC_TABLE}}
        new_para = doc.add_paragraph()
        new_para.text = "{{COMPLEX_DYNAMIC_TABLE}}"
        
        # Insert the paragraph before the table in the XML tree
        tbl_element.addprevious(new_para._element)
        
        # Remove table from XML parent
        parent.remove(tbl_element)
        print("✔ Table successfully replaced with {{COMPLEX_DYNAMIC_TABLE}} placeholder paragraph.")
        
    doc.save(dst_path)
    print(f"✔ Template master generated and saved to: {dst_path}")

if __name__ == '__main__':
    convert_document()
