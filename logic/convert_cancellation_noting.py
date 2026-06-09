import os
import docx

def convert_cancellation_document():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    src_path = os.path.join(base_dir, 'masters', 'For_NominationCancellation_Noting.docx')
    dst_path = os.path.join(base_dir, 'masters', 'Cancellation_Noting_Master.docx')
    
    if not os.path.exists(src_path):
        print(f"Error: {src_path} does not exist!")
        return
        
    doc = docx.Document(src_path)
    
    # ── Paragraph 0: Reference number & Date ──
    p0 = doc.paragraphs[0]
    if len(p0.runs) >= 4:
        p0.runs[2].text = " टीबीआरएल/{{REF_NO}} "
        p0.runs[3].text = "                                                       दिनांक: {{REF_DATE}}"
        
    # ── Paragraph 2: Subject Line ──
    p2 = doc.paragraphs[2]
    if len(p2.runs) >= 8:
        p2.runs[2].text = " “{{SUBJECT_HINDI}}” हेतु कैसिलेशन नामांकन"
        p2.runs[3].text = ""
        p2.runs[4].text = ""
        p2.runs[5].text = ""
        p2.runs[6].text = "।/"
        p2.runs[7].text = "Cancellation Nomination of {{SUBJECT_ENGLISH}}"
        for r in p2.runs[8:]:
            r.text = ""

    # ── Paragraph 4: Reference ──
    p4 = doc.paragraphs[4]
    if len(p4.runs) >= 4:
        p4.runs[1].text = " {{REFERENCE_TEXT}}"
        p4.runs[3].text = "दिनांक: {{REF_MAIL_DATE}}                  ……. ..(1)"

    # ── Paragraph 5: ION Reference ──
    p5 = doc.paragraphs[5]
    if len(p5.runs) >= 1:
        p5.runs[0].text = "            {{ION_REF_SOURCE}}             दिनांक: {{ION_REF_DATE}}               ……. ..(2)"

    # ── Paragraph 7: Intro Text ──
    p7 = doc.paragraphs[7]
    if len(p7.runs) >= 2:
        p7.runs[1].text = " उपरोक्त संलग्न संदर्भ सं. (1), के अनुसार {{COURSE_DATES}} के दौरान {{ORG_INSTITUTE}} द्वारा, {{LAB_NAME}} में \"{{COURSE_TITLE}}\" {{COURSE_TYPE}} में भाग लेने के लिये निम्नलिखित ऑफिसर का नामांकन किया गया था।"
        for r in p7.runs[2:]:
            r.text = ""

    # ── Paragraph 8: Cancellation Reason Text ──
    p8 = doc.paragraphs[8]
    if len(p8.runs) >= 2:
        p8.runs[1].text = "उपरोक्त संलग्न संदर्भ सं. (2), के अनुसार {{CANCEL_NOMINEE_NAME}} का {{CANCEL_GROUP_NAME}} ग्रुप से आई.ओ.एन प्राप्त हुआ है। जिसमे इन्होने {{CANCEL_REASON}} होने का कारण बताया है। जिसके कारण वे इस {{COURSE_TYPE_SHORT}} मे भाग नहीं ले पाएंगे।"
        for r in p8.runs[2:]:
            r.text = ""

    # ── Signatories Paragraphs ──
    p18 = doc.paragraphs[18]
    for run in p18.runs:
        if "राकेश कुमार" in run.text:
            run.text = run.text.replace("राकेश कुमार", "{{SIGNATORY_1_NAME}}")
            
    p19 = doc.paragraphs[19]
    for run in p19.runs:
        if "वैज्ञानिक ‘एफ’(मा.सं.वि.वि)" in run.text:
            run.text = run.text.replace("वैज्ञानिक ‘एफ’(मा.सं.वि.वि)", "{{SIGNATORY_1_DESIG}}")
            
    p21 = doc.paragraphs[21]
    for run in p21.runs:
        if "डॉ. अनिल  खुराना, वैज्ञानिक 'जी'" in run.text:
            run.text = run.text.replace("डॉ. अनिल  खुराना, वैज्ञानिक 'जी'", "{{SIGNATORY_2_NAME}}")
        elif "डॉ. अनिल" in run.text:
            run.text = run.text.replace(run.text, "  {{SIGNATORY_2_NAME}}")
            
    p22 = doc.paragraphs[22]
    for run in p22.runs:
        if "ग्रुप निदेशक(मा.सं.वि.वि.)" in run.text:
            run.text = run.text.replace("ग्रुप निदेशक(मा.सं.वि.वि.)", "{{SIGNATORY_2_DESIG}}")
        elif "ग्रुप निदेशक" in run.text:
            run.text = run.text.replace(run.text, "       {{SIGNATORY_2_DESIG}}")

    # ── Table replacement ──
    if len(doc.tables) > 0:
        table_to_replace = doc.tables[0]
        tbl_element = table_to_replace._tbl
        parent = tbl_element.getparent()
        
        new_para = doc.add_paragraph()
        new_para.text = "{{COMPLEX_DYNAMIC_TABLE}}"
        
        tbl_element.addprevious(new_para._element)
        parent.remove(tbl_element)
        print("✔ Table successfully replaced with {{COMPLEX_DYNAMIC_TABLE}}.")

    doc.save(dst_path)
    print(f"✔ Formatting preserved template master generated at: {dst_path}")

if __name__ == '__main__':
    convert_cancellation_document()
