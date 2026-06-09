import os
import docx

def convert_lecture_document():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    src_path = os.path.join(base_dir, 'masters', 'For_Lecture_Noting.docx')
    dst_path = os.path.join(base_dir, 'masters', 'Lecture_Noting_Master.docx')
    
    if not os.path.exists(src_path):
        print(f"Error: {src_path} does not exist!")
        return
        
    doc = docx.Document(src_path)
    
    # ── Paragraph 0: Reference number & Date ──
    p0 = doc.paragraphs[0]
    if len(p0.runs) >= 3:
        p0.runs[1].text = " {{LAB_NAME}}/{{REF_NO}} "
        p0.runs[2].text = p0.runs[2].text.replace("2025", "{{CURRENT_YEAR}}")
        
    # ── Paragraph 2: Subject Line ──
    p2 = doc.paragraphs[2]
    if len(p2.runs) >= 3:
        p2.runs[2].text = " {{SUBJECT_HINDI}} / {{SUBJECT_ENGLISH}}"
        for r in p2.runs[3:]:
            r.text = ""
        
    # ── Paragraph 4: Reference ──
    p4 = doc.paragraphs[4]
    if len(p4.runs) >= 2:
        p4.runs[1].text = " {{REFERENCE_TEXT}} ..(1)"
        for r in p4.runs[2:]:
            r.text = ""

    # ── Paragraph 7: Intro Text ──
    p7 = doc.paragraphs[7]
    if len(p7.runs) >= 1:
        p7.runs[0].text = "2. उपरोक्त संलग्न संदर्भ सं. (1), के अनुसार {{COURSE_DATES}} के दौरान {{ORG_INSTITUTE}} द्वारा, {{LAB_NAME}} में {{COURSE_TITLE}} विषय पर {{COURSE_TYPE}} आयोजित किया जा रहा है। इस कोर्स में {{LECTURE_TITLE}} विषय पर व्याख्यान देने के लिए {{LAB_NAME}} द्वारा आग्रह किये गये {{OFFICIAL_WORD}} / {{GROUP_NAME}} ग्रुप द्वारा प्राप्त {{NAMANKAN_WORD}} का विवरण निम्नलिखित है।"
        for r in p7.runs[1:]:
            r.text = ""

    # ── Paragraph 9: Request Text ──
    p9 = doc.paragraphs[9]
    if len(p9.runs) >= 1:
        p9.runs[0].text = "4.आपसे निवेदन है कि प्राप्त {{NAMANKAN_WORD}} से या किसी अन्य उपयुक्त ऑफिशियल को उपरोक्त विषय में व्याख्यान देने/भाग लेने के लिए नामांकित करें।"
        for r in p9.runs[1:]:
            r.text = ""

    # ── Signatories Paragraphs ──
    p15 = doc.paragraphs[15]
    for run in p15.runs:
        if "राकेश कुमार" in run.text:
            run.text = run.text.replace("राकेश कुमार", "{{SIGNATORY_1_NAME}}")
            
    p16 = doc.paragraphs[16]
    for run in p16.runs:
        if "वैज्ञानिक ‘एफ’(मा.सं.वि.वि)" in run.text:
            run.text = run.text.replace("वैज्ञानिक ‘एफ’(मा.सं.वि.वि)", "{{SIGNATORY_1_DESIG}}")
            
    p19 = doc.paragraphs[19]
    for run in p19.runs:
        if "डॉ. अनिल  खुराना, वैज्ञानिक 'जी'" in run.text:
            run.text = run.text.replace("डॉ. अनिल  खुराना, वैज्ञानिक 'जी'", "{{SIGNATORY_2_NAME}}")
        elif "डॉ. अनिल" in run.text:
            run.text = run.text.replace(run.text, "  {{SIGNATORY_2_NAME}}")
            
    p20 = doc.paragraphs[20]
    for run in p20.runs:
        if "ग्रुप निदेशक(मा.सं.वि.वि.)" in run.text:
            run.text = run.text.replace("ग्रुप निदेशक(मा.सं.वि.वि.)", "{{SIGNATORY_2_DESIG}}")
        elif "ग्रुप निदेशक" in run.text:
            run.text = run.text.replace(run.text, "       {{SIGNATORY_2_DESIG}}")

    # ── Table replacement ──
    if len(doc.tables) > 0:
        table_to_replace = doc.tables[0]
        tbl_element = table_to_replace._tbl
        parent = tbl_element.getparent()
        
        new_para = doc.add_paragraph()
        new_para.text = "{{COMPLEX_DYNAMIC_TABLE}}"
        
        tbl_element.addprevious(new_para._element)
        parent.remove(tbl_element)
        print("✔ Table successfully replaced with {{COMPLEX_DYNAMIC_TABLE}}.")

    doc.save(dst_path)
    print(f"✔ Formatting preserved template master generated at: {dst_path}")

if __name__ == '__main__':
    convert_lecture_document()
