import os
import docx

def convert_dgmss_document():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    src_path = os.path.join(base_dir, 'masters', 'For_DGMSS_Noting.docx')
    dst_path = os.path.join(base_dir, 'masters', 'DGMSS_Noting_Master.docx')
    
    if not os.path.exists(src_path):
        print(f"Error: {src_path} does not exist!")
        return
        
    doc = docx.Document(src_path)
    
    # ── Paragraph 0: Reference number & Date ──
    p0 = doc.paragraphs[0]
    if len(p0.runs) >= 1:
        p0.runs[0].text = "संदर्भ सं. {{LAB_NAME}}/{{REF_NO}}                                                  दिनांक: {{REF_DATE}}"
        
    # ── Paragraph 2: Subject Line (Hindi) ──
    p2 = doc.paragraphs[2]
    if len(p2.runs) >= 2:
        p2.runs[1].text = "/Sub : {{SUBJECT_HINDI}}"
        
    # ── Paragraph 3: Subject Line (English) ──
    p3 = doc.paragraphs[3]
    if len(p3.runs) >= 1:
        p3.runs[0].text = "            {{SUBJECT_ENGLISH}}"

    # ── Paragraph 4: Reference ──
    p4 = doc.paragraphs[4]
    if len(p4.runs) >= 1:
        p4.runs[0].text = "संदर्भ सं: {{REFERENCE_TEXT}} .....(1)"

    # ── Paragraph 5: Intro Text (Bilingual) ──
    p5 = doc.paragraphs[5]
    if len(p5.runs) >= 2:
        p5.runs[0].text = "2.उपरोक्त संदर्भ सं.(1) के अनुसार {{ORG_INSTITUTE}} द्वारा {{COURSE_DATES}} को “{{COURSE_TITLE}}” {{COURSE_TYPE}} आयोजित किया जा रहा है। इस कोर्स में भाग लेने के लिए {{GROUP_NAME}} ग्रुप से निम्नलिखित {{OFFICIAL_WORD}} {{KA_KE}} {{NAMANKAN_WORD}} प्राप्त हुआ है। आपसे विनम्र निवेदन है कि प्राप्त {{NAMANKAN_WORD}} से या किसी अन्य उपयुक्त ऑफिसियल्स को इस कोर्स में भाग लेने के लिए नामांकित करें। विस्तृत विवरण निम्नलिखित है।"
        p5.runs[1].text = "(With reference to letter no. (1), {{ORG_INSTITUTE}} is organizing {{COURSE_TITLE}} {{COURSE_TYPE}} during {{COURSE_DATES}} at Texas. Following nomination is received from {{GROUP_NAME}} group to attend this course. Details are given below):-"
        for r in p5.runs[2:]:
            r.text = ""

    # ── Paragraph 6: Request Text ──
    p6 = doc.paragraphs[6]
    if len(p6.runs) >= 1:
        p6.runs[0].text = "3.आपसे विनम्र निवेदन है कि प्राप्त {{NAMANKAN_WORD}} में से या उपयुक्त ऑफिसियल को इस {{COURSE_TYPE}} में भाग लेने के लिये नामांकित करें। (It is requested to kindly nominate received nomination or suitable Official to attend this {{COURSE_TYPE}}.)"
        for r in p6.runs[1:]:
            r.text = ""

    # ── Signatories Paragraphs ──
    p12 = doc.paragraphs[12]
    for run in p12.runs:
        if "राकेश कुमार" in run.text:
            run.text = run.text.replace("राकेश कुमार", "{{SIGNATORY_1_NAME}}")
            
    p13 = doc.paragraphs[13]
    for run in p13.runs:
        if "वैज्ञानिक ‘एफ’(मा.सं.वि.वि)" in run.text:
            run.text = run.text.replace("वैज्ञानिक ‘एफ’(मा.सं.वि.वि)", "{{SIGNATORY_1_DESIG}}")
            
    p14 = doc.paragraphs[14]
    for run in p14.runs:
        if "डॉ. अनिल  खुराना, वैज्ञानिक 'जी'" in run.text:
            run.text = run.text.replace("डॉ. अनिल  खुराना, वैज्ञानिक 'जी'", "{{SIGNATORY_2_NAME}}")
        elif "डॉ. अनिल" in run.text:
            run.text = run.text.replace(run.text, "  {{SIGNATORY_2_NAME}}")
            
    p15 = doc.paragraphs[15]
    for run in p15.runs:
        if "ग्रुप निदेशक(मा.सं.वि.वि.)" in run.text:
            run.text = run.text.replace("ग्रुप निदेशक(मा.सं.वि.वि.)", "{{SIGNATORY_2_DESIG}}")
        elif "ग्रुप निदेशक" in run.text:
            run.text = run.text.replace(run.text, "       {{SIGNATORY_2_DESIG}}")

    # ── Table replacement ──
    if len(doc.tables) > 0:
        table_to_replace = doc.tables[0]
        tbl_element = table_to_replace._tbl
        parent = tbl_element.getparent()
        
        new_para = doc.add_paragraph()
        new_para.text = "{{COMPLEX_DYNAMIC_TABLE}}"
        
        tbl_element.addprevious(new_para._element)
        parent.remove(tbl_element)
        print("✔ Table successfully replaced with {{COMPLEX_DYNAMIC_TABLE}}.")

    doc.save(dst_path)
    print(f"✔ Formatting preserved template master generated at: {dst_path}")

if __name__ == '__main__':
    convert_dgmss_document()
