from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import os
import json
import shutil
from datetime import datetime
from logic.docx_utils import get_safe, parse_date_safe, replace_placeholders_in_paragraph, has_unreplaced_placeholders

OUTPUT_DIR = "generated_notices"




def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    """Sets compact cell padding (in dxas) for tighter table density."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def apply_table_formatting(table):
    """Enforces cantSplit on all rows and tblHeader on the first row to repeat headers and prevent row splitting."""
    if len(table.rows) > 0:
        trPr = table.rows[0]._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def process_reference_paragraph(doc, data):
    """Rebuilds the reference paragraph to format multiple references dynamically."""
    references = data.get("references", [])
    if not references:
        # Fallback to single reference
        ref_no = data.get("ref_no", "")
        lab_name = data.get('lab_name', 'टीबीआरएल')
        source_val = ref_no
        if ref_no and not ref_no.startswith(lab_name):
            source_val = f"{lab_name}/{ref_no}"
        references = [{
            "source": source_val,
            "date": data.get("ref_date", "")
        }]

    # Rebuild the paragraph containing {{REFERENCE_TEXT}}
    ref_para = None
    for para in doc.paragraphs:
        if "{{REFERENCE_TEXT}}" in para.text:
            ref_para = para
            break

    if ref_para:
        # Save formatting from the first run
        font_name = None
        font_size = Pt(10)
        font_bold = True
        if len(ref_para.runs) > 0:
            if ref_para.runs[0].font.name:
                font_name = ref_para.runs[0].font.name
            if ref_para.runs[0].font.size:
                font_size = ref_para.runs[0].font.size
            if ref_para.runs[0].font.bold is not None:
                font_bold = ref_para.runs[0].font.bold

        # Clear runs
        ref_para.text = ""
        
        # Build references list
        for idx, ref in enumerate(references, start=1):
            source = ref.get("source", "").strip()
            date = ref.get("date", "").strip()
            if idx == 1:
                line_text = f"संदर्भ : {source}    दिनांक: {date}    ....... ..({idx})"
            else:
                line_text = f"        {source}    दिनांक: {date}    ....... ..({idx})"
            
            run = ref_para.add_run(line_text)
            run.font.bold = font_bold
            if font_name:
                run.font.name = font_name
            run.font.size = font_size
            
            if idx < len(references):
                ref_para.add_run("\n")

    # Clear ION reference paragraph in Cancellation Noting since it's consolidated
    for para in doc.paragraphs:
        if "{{ION_REF_SOURCE}}" in para.text or "{{ION_REF_DATE}}" in para.text:
            para.text = ""


def build_nominees_table(doc, table_index, columns, nominees):
    """Generates the nominees table with auto S.No and a permanent blank row."""
    # Increase rows by +2 (1 header, len(nominees) data rows, 1 blank row)
    table = doc.add_table(rows=len(nominees) + 2, cols=len(columns))
    set_table_borders(table, color="A0AEC0", sz="4")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Draw dynamic headers
    hdr_cells = table.rows[0].cells
    for col_idx, col_name in enumerate(columns):
        cell = hdr_cells[col_idx]
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        run = p.add_run(col_name)
        run.font.bold = True
        run.font.size = Pt(9)
        
    sn_aliases = ["s.n.", "s.no.", "sl.no.", "cr.sn.", "क्र.सं.", "क्र.स", "क्र.सं", "क्र.स."]
    
    # Draw dynamic data rows
    for row_idx, nominee_data in enumerate(nominees, start=1):
        row_cells = table.rows[row_idx].cells
        
        for col_idx, col_name in enumerate(columns):
            cell = row_cells[col_idx]
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            # Clean col parts (splitting by '/') to check for serial number keywords bilingually
            clean_parts = [p.strip().lower().replace(" ", "").replace(".", "") for p in col_name.split("/")]
            is_sn = any(p in ["sn", "sno", "slno", "crsn", "क्रसं", "क्रस"] for p in clean_parts)
            if is_sn:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_value = f"{row_idx}."
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell_value = str(nominee_data.get(col_name, ""))
            
            # Auto-shrink font if table has many columns to fit A4 page
            font_size = 8.0 if len(columns) > 7 else (8.5 if len(columns) > 5 else 9.5)
            
            run = p.add_run(cell_value)
            run.font.size = Pt(font_size)
            
    # Draw permanent blank row at the end
    blank_row_idx = len(nominees) + 1
    row_cells = table.rows[blank_row_idx].cells
    for col_idx, col_name in enumerate(columns):
        cell = row_cells[col_idx]
        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        # Clean col parts (splitting by '/') to check for serial number keywords bilingually
        clean_parts = [p.strip().lower().replace(" ", "").replace(".", "") for p in col_name.split("/")]
        is_sn = any(p in ["sn", "sno", "slno", "crsn", "क्रसं", "क्रस"] for p in clean_parts)
        if is_sn:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_value = f"{blank_row_idx}."
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell_value = ""
            
        font_size = 8.0 if len(columns) > 7 else (8.5 if len(columns) > 5 else 9.5)
        run = p.add_run(cell_value)
        run.font.size = Pt(font_size)
        
    apply_table_formatting(table)
    # Move table to correct position in document XML
    doc.paragraphs[table_index]._element.addnext(table._element)


def get_course_dates_str(start_date_str, end_date_str):
    start_date = parse_date_safe(start_date_str)
    end_date = parse_date_safe(end_date_str)
    days_count = (end_date - start_date).days + 1
    
    hindi_months = {
        1: "जनवरी", 2: "फ़रवरी", 3: "मार्च", 4: "अप्रैल",
        5: "मई", 6: "जून", 7: "जुलाई", 8: "अगस्त",
        9: "सितंबर", 10: "अक्टूबर", 11: "नवंबर", 12: "दिसंबर"
    }
    start_day = start_date.day
    end_day = end_date.day
    start_month = hindi_months[start_date.month]
    end_month = hindi_months[end_date.month]
    start_year = start_date.year
    end_year = end_date.year
    
    if start_month == end_month and start_year == end_year:
        course_dates_str = f"{start_day}-{end_day} {start_month} {start_year}"
    elif start_year == end_year:
        course_dates_str = f"{start_day} {start_month} से {end_day} {end_month} {start_year}"
    else:
        course_dates_str = f"{start_day} {start_month} {start_year} से {end_day} {end_month} {end_year}"
        
    return course_dates_str, days_count


def generate_tbrl_noting(data):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    master_path = os.path.join("masters", "TBRL_Noting_Master.docx")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    group = data.get('group_name', 'General').replace("/", "-")
    filename = f"TBRL_Noting_{group}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    
    shutil.copy2(master_path, output_path)
    doc = Document(output_path)
    
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    course_dates_str, days_count = get_course_dates_str(data['start_date'], data['end_date'])
    
    num_nominees = len(data.get('nominees', []))
    if num_nominees <= 1:
        off_word, ka_ke, hua_hue, nam_word = "अधिकारी", "का", "हुआ", "नामांकन"
    else:
        off_word, ka_ke, hua_hue, nam_word = "अधिकारियों", "के", "हुए", "नामांकनों"
        
    ref_no = data.get("ref_no", "")
    lab_name = data.get("lab_name", "टीबीआरएल")
    if ref_no.startswith(f"{lab_name}/"):
        ref_no = ref_no[len(lab_name)+1:]

    references = data.get("references", [])
    ref_source = ""
    ref_mail_date = ""
    if len(references) > 1:
        ref_source = references[1].get("source", "")
        ref_mail_date = references[1].get("date", "")
    else:
        ref_source = data.get("reference_text", "")
        ref_mail_date = data.get("ref_date", "")
        
    placeholders = {
        "{{REF_NO}}": ref_no,
        "{{LAB_NAME}}": lab_name,
        "{{CURRENT_YEAR}}": str(datetime.now().year),
        "{{SUBJECT_HINDI}}": data.get("subject_hindi", ""),
        "{{SUBJECT_ENGLISH}}": data.get("subject_english", ""),
        "{{COURSE_TYPE}}": data.get("course_type", ""),
        "{{REFERENCE_TEXT}}": ref_source,
        "{{REF_DATE}}": data.get("ref_date", ""),
        "{{COURSE_DATES}}": course_dates_str,
        "{{START_DATE}}": parse_date_safe(get_safe(data, 'start_date')).strftime('%d %B %Y'),
        "{{END_DATE}}": parse_date_safe(get_safe(data, 'end_date')).strftime('%d %B %Y'),
        "{{DAYS_COUNT}}": str(days_count),
        "{{ORG_INSTITUTE}}": data.get("org_institute", ""),
        "{{COURSE_TITLE}}": data.get("course_title", ""),
        "{{GROUP_NAME}}": f"{data.get('group_name', '')} Group",
        "{{OFFICIAL_WORD}}": off_word,
        "{{KA_KE}}": ka_ke,
        "{{HUA_HUE}}": hua_hue,
        "{{NAMANKAN_WORD}}": nam_word,
        "{{SIGNATORY_1_NAME}}": get_safe(data, "sig1_name", ""),
        "{{SIGNATORY_1_DESIG}}": get_safe(data, "sig1_desig", ""),
        "{{SIGNATORY_2_NAME}}": get_safe(data, "sig2_name", ""),
        "{{SIGNATORY_2_DESIG}}": get_safe(data, "sig2_desig", ""),
        "संदर्भ सं. (1)": "संदर्भ सं. (2)",
        "संदर्भ सं. (2)": "संदर्भ सं. (3)",
    }

    process_reference_paragraph(doc, data)

    table_index = None
    for i, para in enumerate(doc.paragraphs):
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.05
        
        for run in para.runs:
            if run.font.size is None or run.font.size > Pt(11):
                run.font.size = Pt(10)
                
        if "{{COMPLEX_DYNAMIC_TABLE}}" in para.text:
            table_index = i
            para.text = "" 
        else:
            replace_placeholders_in_paragraph(para, placeholders)

    if table_index is not None and data.get('columns') and data.get('nominees'):
        build_nominees_table(doc, table_index, data['columns'], data['nominees'])

    if has_unreplaced_placeholders(doc):
        print("WARNING: Unreplaced placeholders detected in generated document!")
    doc.save(output_path)
    return output_path, filename


def generate_lecture_noting(data):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    master_path = os.path.join("masters", "Lecture_Noting_Master.docx")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    group = data.get('group_name', 'General').replace("/", "-")
    filename = f"Lecture_Noting_{group}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    
    shutil.copy2(master_path, output_path)
    doc = Document(output_path)
    
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    course_dates_str, days_count = get_course_dates_str(data['start_date'], data['end_date'])
    
    num_nominees = len(data.get('nominees', []))
    if num_nominees <= 1:
        off_word, ka_ke, hua_hue, nam_word = "अधिकारी", "का", "हुआ", "नामांकन"
    else:
        off_word, ka_ke, hua_hue, nam_word = "अधिकारियों", "के", "हुए", "नामांकनों"
        
    ref_no = data.get("ref_no", "")
    lab_name = data.get("lab_name", "टीबीआरएल")
    if ref_no.startswith(f"{lab_name}/"):
        ref_no = ref_no[len(lab_name)+1:]

    references = data.get("references", [])
    ref_source = ""
    ref_mail_date = ""
    if len(references) > 1:
        ref_source = references[1].get("source", "")
        ref_mail_date = references[1].get("date", "")
    else:
        ref_source = data.get("reference_text", "")
        ref_mail_date = data.get("ref_date", "")
        
    placeholders = {
        "{{REF_NO}}": ref_no,
        "{{LAB_NAME}}": lab_name,
        "{{CURRENT_YEAR}}": str(datetime.now().year),
        "{{SUBJECT_HINDI}}": data.get("subject_hindi", ""),
        "{{SUBJECT_ENGLISH}}": data.get("subject_english", ""),
        "{{COURSE_TYPE}}": data.get("course_type", ""),
        "{{REFERENCE_TEXT}}": ref_source,
        "{{REF_DATE}}": data.get("ref_date", ""),
        "{{COURSE_DATES}}": course_dates_str,
        "{{START_DATE}}": parse_date_safe(get_safe(data, 'start_date')).strftime('%d %B %Y'),
        "{{END_DATE}}": parse_date_safe(get_safe(data, 'end_date')).strftime('%d %B %Y'),
        "{{DAYS_COUNT}}": str(days_count),
        "{{ORG_INSTITUTE}}": data.get("org_institute", ""),
        "{{COURSE_TITLE}}": data.get("course_title", ""),
        "{{LECTURE_TITLE}}": data.get("lecture_title", ""),
        "{{GROUP_NAME}}": f"{data.get('group_name', '')} Group",
        "{{OFFICIAL_WORD}}": off_word,
        "{{KA_KE}}": ka_ke,
        "{{HUA_HUE}}": hua_hue,
        "{{NAMANKAN_WORD}}": nam_word,
        "{{SIGNATORY_1_NAME}}": get_safe(data, "sig1_name", ""),
        "{{SIGNATORY_1_DESIG}}": get_safe(data, "sig1_desig", ""),
        "{{SIGNATORY_2_NAME}}": get_safe(data, "sig2_name", ""),
        "{{SIGNATORY_2_DESIG}}": get_safe(data, "sig2_desig", ""),
        "संदर्भ सं. (1)": "संदर्भ सं. (2)",
        "संदर्भ सं. (2)": "संदर्भ सं. (3)",
    }

    process_reference_paragraph(doc, data)

    table_index = None
    for i, para in enumerate(doc.paragraphs):
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.05
        
        for run in para.runs:
            if run.font.size is None or run.font.size > Pt(11):
                run.font.size = Pt(10)
                
        if "{{COMPLEX_DYNAMIC_TABLE}}" in para.text:
            table_index = i
            para.text = "" 
        else:
            replace_placeholders_in_paragraph(para, placeholders)

    if table_index is not None and data.get('columns') and data.get('nominees'):
        build_nominees_table(doc, table_index, data['columns'], data['nominees'])

    if has_unreplaced_placeholders(doc):
        print("WARNING: Unreplaced placeholders detected in generated document!")
    doc.save(output_path)
    return output_path, filename


def generate_dgmss_noting(data):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    master_path = os.path.join("masters", "DGMSS_Noting_Master.docx")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    group = data.get('group_name', 'General').replace("/", "-")
    filename = f"DGMSS_Noting_{group}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    
    shutil.copy2(master_path, output_path)
    doc = Document(output_path)
    
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    course_dates_str, days_count = get_course_dates_str(data['start_date'], data['end_date'])
    
    num_nominees = len(data.get('nominees', []))
    if num_nominees <= 1:
        off_word, ka_ke, hua_hue, nam_word = "अधिकारी", "का", "हुआ", "नामांकन"
    else:
        off_word, ka_ke, hua_hue, nam_word = "अधिकारियों", "के", "हुए", "नामांकनों"
        
    ref_no = data.get("ref_no", "")
    lab_name = data.get("lab_name", "टीबीआरएल")
    if ref_no.startswith(f"{lab_name}/"):
        ref_no = ref_no[len(lab_name)+1:]

    references = data.get("references", [])
    ref_source = ""
    ref_mail_date = ""
    if len(references) > 1:
        ref_source = references[1].get("source", "")
        ref_mail_date = references[1].get("date", "")
    else:
        ref_source = data.get("reference_text", "")
        ref_mail_date = data.get("ref_date", "")
        
    placeholders = {
        "{{REF_NO}}": ref_no,
        "{{LAB_NAME}}": lab_name,
        "{{CURRENT_YEAR}}": str(datetime.now().year),
        "{{SUBJECT_HINDI}}": data.get("subject_hindi", ""),
        "{{SUBJECT_ENGLISH}}": data.get("subject_english", ""),
        "{{COURSE_TYPE}}": data.get("course_type", ""),
        "{{REFERENCE_TEXT}}": ref_source,
        "{{REF_DATE}}": data.get("ref_date", ""),
        "{{COURSE_DATES}}": course_dates_str,
        "{{START_DATE}}": parse_date_safe(get_safe(data, 'start_date')).strftime('%d %B %Y'),
        "{{END_DATE}}": parse_date_safe(get_safe(data, 'end_date')).strftime('%d %B %Y'),
        "{{DAYS_COUNT}}": str(days_count),
        "{{ORG_INSTITUTE}}": data.get("org_institute", ""),
        "{{COURSE_TITLE}}": data.get("course_title", ""),
        "{{GROUP_NAME}}": f"{data.get('group_name', '')} Group",
        "{{OFFICIAL_WORD}}": off_word,
        "{{KA_KE}}": ka_ke,
        "{{HUA_HUE}}": hua_hue,
        "{{NAMANKAN_WORD}}": nam_word,
        "{{SIGNATORY_1_NAME}}": get_safe(data, "sig1_name", ""),
        "{{SIGNATORY_1_DESIG}}": get_safe(data, "sig1_desig", ""),
        "{{SIGNATORY_2_NAME}}": get_safe(data, "sig2_name", ""),
        "{{SIGNATORY_2_DESIG}}": get_safe(data, "sig2_desig", ""),
        "संदर्भ सं. (1)": "संदर्भ सं. (2)",
        "संदर्भ सं. (2)": "संदर्भ सं. (3)",
    }

    process_reference_paragraph(doc, data)

    table_index = None
    for i, para in enumerate(doc.paragraphs):
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.05
        
        for run in para.runs:
            if run.font.size is None or run.font.size > Pt(11):
                run.font.size = Pt(10)
                
        if "{{COMPLEX_DYNAMIC_TABLE}}" in para.text:
            table_index = i
            para.text = "" 
        else:
            replace_placeholders_in_paragraph(para, placeholders)

    if table_index is not None and data.get('columns') and data.get('nominees'):
        build_nominees_table(doc, table_index, data['columns'], data['nominees'])

    if has_unreplaced_placeholders(doc):
        print("WARNING: Unreplaced placeholders detected in generated document!")
    doc.save(output_path)
    return output_path, filename


def generate_fee_noting(data):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    master_path = os.path.join("masters", "Fee_Noting_Master.docx")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    group = data.get('group_name', 'General').replace("/", "-")
    filename = f"Fee_Noting_{group}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    
    shutil.copy2(master_path, output_path)
    doc = Document(output_path)
    
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    course_dates_str, days_count = get_course_dates_str(data['start_date'], data['end_date'])
    
    num_nominees = len(data.get('nominees', []))
    if num_nominees <= 1:
        off_word, ka_ke, hua_hue, nam_word = "अधिकारी", "का", "नामांकन", "विवरण"
    else:
        off_word, ka_ke, hua_hue, nam_word = "अधिकारियों", "के", "नामांकनों", "विवरण"
        
    ref_no = data.get("ref_no", "")
    lab_name = data.get("lab_name", "टीबीआरएल")
    if ref_no.startswith(f"{lab_name}/"):
        ref_no = ref_no[len(lab_name)+1:]

    references = data.get("references", [])
    ref_source = ""
    ref_mail_date = ""
    if len(references) > 1:
        ref_source = references[1].get("source", "")
        ref_mail_date = references[1].get("date", "")
    else:
        ref_source = data.get("reference_text", "")
        ref_mail_date = data.get("ref_mail_date", "")

    placeholders = {
        "{{REF_NO}}": ref_no,
        "{{LAB_NAME}}": lab_name,
        "{{CURRENT_YEAR}}": str(datetime.now().year),
        "{{SUBJECT_HINDI}}": data.get("subject_hindi", ""),
        "{{SUBJECT_ENGLISH}}": data.get("subject_english", ""),
        "{{COURSE_TYPE}}": data.get("course_type", ""),
        "{{REFERENCE_TEXT}}": ref_source,
        "{{REF_DATE}}": data.get("ref_date", ""),
        "{{REF_MAIL_DATE}}": ref_mail_date,
        "{{COURSE_DATES}}": course_dates_str,
        "{{START_DATE}}": parse_date_safe(get_safe(data, 'start_date')).strftime('%d %B %Y'),
        "{{END_DATE}}": parse_date_safe(get_safe(data, 'end_date')).strftime('%d %B %Y'),
        "{{DAYS_COUNT}}": str(days_count),
        "{{ORG_INSTITUTE}}": data.get("org_institute", ""),
        "{{COURSE_TITLE}}": data.get("course_title", ""),
        "{{GROUP_NAME}}": f"{data.get('group_name', '')} Group",
        "{{OFFICIAL_WORD}}": off_word,
        "{{KA_KE}}": ka_ke,
        "{{HUA_HUE}}": hua_hue,
        "{{NAMANKAN_WORD}}": nam_word,
        "{{SIGNATORY_1_NAME}}": get_safe(data, "sig1_name", ""),
        "{{SIGNATORY_1_DESIG}}": get_safe(data, "sig1_desig", ""),
        "{{SIGNATORY_2_NAME}}": get_safe(data, "sig2_name", ""),
        "{{SIGNATORY_2_DESIG}}": get_safe(data, "sig2_desig", ""),
        "संदर्भ सं. (1)": "संदर्भ सं. (2)",
        "संदर्भ सं. (2)": "संदर्भ सं. (3)",
    }

    process_reference_paragraph(doc, data)

    table_index = None
    for i, para in enumerate(doc.paragraphs):
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.05
        
        for run in para.runs:
            if run.font.size is None or run.font.size > Pt(11):
                run.font.size = Pt(10)
                
        if "{{COMPLEX_DYNAMIC_TABLE}}" in para.text:
            table_index = i
            para.text = "" 
        else:
            replace_placeholders_in_paragraph(para, placeholders)

    if table_index is not None and data.get('columns') and data.get('nominees'):
        build_nominees_table(doc, table_index, data['columns'], data['nominees'])

    if has_unreplaced_placeholders(doc):
        print("WARNING: Unreplaced placeholders detected in generated document!")
    doc.save(output_path)
    return output_path, filename


def generate_cancellation_noting(data):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    master_path = os.path.join("masters", "Cancellation_Noting_Master.docx")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    group = data.get('group_name', 'General').replace("/", "-")
    filename = f"Cancellation_Noting_{group}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    
    shutil.copy2(master_path, output_path)
    doc = Document(output_path)
    
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    course_dates_str, days_count = get_course_dates_str(data['start_date'], data['end_date'])
    
    ref_no = data.get("ref_no", "")
    lab_name = data.get("lab_name", "टीबीआरएल")
    if ref_no.startswith(f"{lab_name}/"):
        ref_no = ref_no[len(lab_name)+1:]

    references = data.get("references", [])
    ref_source = ""
    ref_mail_date = ""
    ion_ref_source = ""
    ion_ref_date = ""
    
    if len(references) > 1:
        ref_source = references[1].get("source", "")
        ref_mail_date = references[1].get("date", "")
    else:
        ref_source = data.get("reference_text", "")
        ref_mail_date = data.get("ref_mail_date", "")
        
    if len(references) > 2:
        ion_ref_source = references[2].get("source", "")
        ion_ref_date = references[2].get("date", "")
    else:
        ion_ref_source = data.get("ion_ref_source", "")
        ion_ref_date = data.get("ion_ref_date", "")

    placeholders = {
        "{{REF_NO}}": ref_no,
        "{{LAB_NAME}}": lab_name,
        "{{CURRENT_YEAR}}": str(datetime.now().year),
        "{{SUBJECT_HINDI}}": data.get("subject_hindi", ""),
        "{{SUBJECT_ENGLISH}}": data.get("subject_english", ""),
        "{{COURSE_TYPE}}": data.get("course_type", ""),
        "{{REFERENCE_TEXT}}": ref_source,
        "{{REF_DATE}}": data.get("ref_date", ""),
        "{{REF_MAIL_DATE}}": ref_mail_date,
        "{{ION_REF_SOURCE}}": ion_ref_source,
        "{{ION_REF_DATE}}": ion_ref_date,
        "{{CANCEL_NOMINEE_NAME}}": data.get("cancel_nominee_name", ""),
        "{{CANCEL_GROUP_NAME}}": data.get("cancel_group_name", ""),
        "{{CANCEL_REASON}}": data.get("cancel_reason", ""),
        "{{COURSE_TYPE_SHORT}}": data.get("course_type_short", ""),
        "{{COURSE_DATES}}": course_dates_str,
        "{{START_DATE}}": parse_date_safe(get_safe(data, 'start_date')).strftime('%d %B %Y'),
        "{{END_DATE}}": parse_date_safe(get_safe(data, 'end_date')).strftime('%d %B %Y'),
        "{{DAYS_COUNT}}": str(days_count),
        "{{ORG_INSTITUTE}}": data.get("org_institute", ""),
        "{{COURSE_TITLE}}": data.get("course_title", ""),
        "{{GROUP_NAME}}": f"{data.get('group_name', '')} Group",
        "{{SIGNATORY_1_NAME}}": get_safe(data, "sig1_name", ""),
        "{{SIGNATORY_1_DESIG}}": get_safe(data, "sig1_desig", ""),
        "{{SIGNATORY_2_NAME}}": get_safe(data, "sig2_name", ""),
        "{{SIGNATORY_2_DESIG}}": get_safe(data, "sig2_desig", ""),
        "संदर्भ सं. (1)": "संदर्भ सं. (2)",
        "संदर्भ सं. (2)": "संदर्भ सं. (3)",
    }

    process_reference_paragraph(doc, data)

    table_index = None
    for i, para in enumerate(doc.paragraphs):
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.05
        
        for run in para.runs:
            if run.font.size is None or run.font.size > Pt(11):
                run.font.size = Pt(10)
                
        if "{{COMPLEX_DYNAMIC_TABLE}}" in para.text:
            table_index = i
            para.text = "" 
        else:
            replace_placeholders_in_paragraph(para, placeholders)

    if table_index is not None and data.get('columns') and data.get('nominees'):
        build_nominees_table(doc, table_index, data['columns'], data['nominees'])

    if has_unreplaced_placeholders(doc):
        print("WARNING: Unreplaced placeholders detected in generated document!")
    doc.save(output_path)
    return output_path, filename


def build_dynamic_fax_table(doc, table_index, columns, rows):
    """Generates a dynamic table for fax sheets without a permanent blank row."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
    set_table_borders(table, color="A0AEC0", sz="4")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Draw dynamic headers
    hdr_cells = table.rows[0].cells
    for col_idx, col_name in enumerate(columns):
        cell = hdr_cells[col_idx]
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        run = p.add_run(col_name)
        run.font.bold = True
        run.font.size = Pt(9.5)
        
    sn_aliases = ["s.n.", "s.no.", "sl.no.", "cr.sn.", "क्र.सं.", "क्र.स", "क्र.सं", "क्र.स."]
    
    # Draw dynamic data rows
    for row_idx, row_data in enumerate(rows, start=1):
        row_cells = table.rows[row_idx].cells
        
        for col_idx, col_name in enumerate(columns):
            cell = row_cells[col_idx]
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            # Clean col parts (splitting by '/') to check for serial number keywords bilingually
            clean_parts = [p.strip().lower().replace(" ", "").replace(".", "") for p in col_name.split("/")]
            is_sn = any(p in ["sn", "sno", "slno", "crsn", "क्रसं", "क्रस"] for p in clean_parts)
            if is_sn:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_value = f"{row_idx}."
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell_value = str(row_data.get(col_name, ""))
            
            font_size = 8.0 if len(columns) > 7 else (8.5 if len(columns) > 5 else 9.5)
            run = p.add_run(cell_value)
            run.font.size = Pt(font_size)
            
    apply_table_formatting(table)
    doc.paragraphs[table_index]._element.addnext(table._element)


def generate_date_amendment_fax(data):
    """Generates the Date Amendment Fax sheet from its master template."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    master_path = os.path.join("masters", "Date_Amendment_Fax.docx")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Date_Amendment_Fax_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    
    shutil.copy2(master_path, output_path)
    doc = Document(output_path)
    
    ref_no = data.get("ref_no", "")
    ref_date = data.get("ref_date", "")
    if not ref_date.strip():
        ref_date = "  /    /2026"
    fax_no = data.get("fax_no", "")
    to_text = data.get("to_text", "")
    
    subject_hindi = data.get("subject_hindi", "")
    subject_english = data.get("subject_english", "")
    ref_text = data.get("ref_text", "")
    
    ref_to_hindi = data.get("ref_to_hindi", "")
    ref_to_eng = data.get("ref_to_eng", "")
    num_courses_hindi = data.get("num_courses_hindi", "तीन")
    num_courses_eng = data.get("num_courses_eng", "three")
    courses_plural_eng = data.get("courses_plural_eng", "Courses")
    amended_courses_hindi = data.get("amended_courses_hindi", "एक")
    amended_courses_eng = data.get("amended_courses_eng", "one")
    
    placeholders = {
        "{{REF_NO}}": ref_no,
        "{{REF_DATE}}": ref_date,
        "{{FAX_NO}}": fax_no,
        "{{SUBJECT_HINDI}}": subject_hindi,
        "{{SUBJECT_ENGLISH}}": subject_english,
        "{{REF_TEXT}}": ref_text,
        "{{REF_TO_HINDI}}": ref_to_hindi,
        "{{REF_TO_ENG}}": ref_to_eng,
        "{{NUM_COURSES_HINDI}}": num_courses_hindi,
        "{{NUM_COURSES_ENG}}": num_courses_eng,
        "{{COURSES_PLURAL_ENG}}": courses_plural_eng,
        "{{AMENDED_COURSES_HINDI}}": amended_courses_hindi,
        "{{AMENDED_COURSES_ENG}}": amended_courses_eng,
        "{{SIG_NAME}}": data.get("sig_name", ""),
        "{{SIG_DESIG}}": data.get("sig_desig", ""),
        "{{FOR_DIRECTOR}}": data.get("for_director", ""),
    }
    
    table_index_for = None
    table_index_read = None
    
    for i, para in enumerate(doc.paragraphs):
        if "{{FOR_TABLE}}" in para.text:
            table_index_for = i
            para.text = ""
        elif "{{READ_TABLE}}" in para.text:
            table_index_read = i
            para.text = ""
        elif "{{TO_TEXT}}" in para.text:
            para.text = ""
            lines = to_text.split("\n")
            for idx, line in enumerate(lines):
                if idx > 0:
                    para.add_run("\n")
                r = para.add_run(line)
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
        else:
            replace_placeholders_in_paragraph(para, placeholders)
            
    # Draw first table ("For")
    if table_index_for is not None and data.get("for_columns") and data.get("for_rows"):
        build_dynamic_fax_table(doc, table_index_for, data["for_columns"], data["for_rows"])
        
    # Draw second table ("Read")
    if table_index_read is not None and data.get("read_columns") and data.get("read_rows"):
        build_dynamic_fax_table(doc, table_index_read, data["read_columns"], data["read_rows"])
        
    if has_unreplaced_placeholders(doc):
        print("WARNING: Unreplaced placeholders detected in generated document!")
    doc.save(output_path)
    return output_path, filename


def generate_mayurpankh_erp_fax(data):
    """Generates the Mayurpankh ERP Fax/Notice sheet from its master template."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    master_path = os.path.join("masters", "Mayurpankh_ERP_Notice.docx")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Mayurpankh_ERP_Notice_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    
    shutil.copy2(master_path, output_path)
    doc = Document(output_path)
    
    ref_no = data.get("ref_no", "")
    ref_date = data.get("ref_date", "")
    if not ref_date.strip():
        ref_date = "  /    /2026"
    fax_no = data.get("fax_no", "")
    to_text = data.get("to_text", "")
    subject_hindi = data.get("subject_hindi", "")
    subject_english = data.get("subject_english", "")
    ref_text = data.get("ref_text", "")
    attn_text = data.get("attn_text", "")
    body_hindi = data.get("body_hindi", "")
    body_english = data.get("body_english", "")
    confirm_text = data.get("confirm_text", "")
    
    placeholders = {
        "{{REF_NO}}": ref_no,
        "{{REF_DATE}}": ref_date,
        "{{FAX_NO}}": fax_no,
        "{{SUBJECT_HINDI}}": subject_hindi,
        "{{SUBJECT_ENGLISH}}": subject_english,
        "{{REF_TEXT}}": ref_text,
        "{{SIG_NAME}}": data.get("sig_name", ""),
        "{{SIG_DESIG}}": data.get("sig_desig", ""),
        "{{FOR_DIRECTOR}}": data.get("for_director", ""),
    }
    
    table_index = None
    
    for i, para in enumerate(doc.paragraphs):
        if "{{NOMINEES_TABLE}}" in para.text:
            table_index = i
            para.text = ""
        elif "{{TO_TEXT}}" in para.text:
            para.text = ""
            lines = to_text.split("\n")
            for idx, line in enumerate(lines):
                if idx > 0:
                    para.add_run("\n")
                r = para.add_run(line)
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
        elif "{{ATTN_TEXT}}" in para.text:
            para.text = ""
            lines = attn_text.split("\n")
            for idx, line in enumerate(lines):
                if idx > 0:
                    para.add_run("\n")
                if idx == 0:
                    lbl = para.add_run("ध्यानार्थ (Kindly Attention): ")
                    lbl.font.name = "Times New Roman"
                    lbl.font.bold = True
                    lbl.font.size = Pt(10)
                    r = para.add_run(line)
                else:
                    prefix = "                                                       "
                    r = para.add_run(prefix + line)
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
        elif "{{BODY_HINDI}}" in para.text:
            para.text = ""
            r = para.add_run(body_hindi)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
        elif "{{BODY_ENGLISH}}" in para.text:
            para.text = ""
            r = para.add_run(body_english)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
        elif "{{CONFIRM_TEXT}}" in para.text:
            para.text = ""
            r = para.add_run(confirm_text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
        else:
            replace_placeholders_in_paragraph(para, placeholders)
            
    # Draw the dynamic table
    # Draw the dynamic table
    if table_index is not None and data.get("columns") and data.get("rows"):
        build_dynamic_fax_table(doc, table_index, data["columns"], data["rows"])
        
    if has_unreplaced_placeholders(doc):
        print("WARNING: Unreplaced placeholders detected in generated document!")
    doc.save(output_path)
    return output_path, filename


def generate_appraisal_proforma(data):
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    template_path = os.path.join(base_dir, "masters", "Appraisal_Proforma_Master.docx")
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at: {template_path}")
        
    doc = Document(template_path)
    
    # Placeholders mapping
    placeholders = {
        "{{STUDENT_NAME_HINDI}}":   get_safe(data, "student_name_hindi", ""),
        "{{FATHER_NAME_HINDI}}":    get_safe(data, "father_name_hindi", ""),
        "{{STUDENT_NAME_ENG}}":     get_safe(data, "student_name_eng", ""),
        "{{FATHER_NAME_ENG}}":      get_safe(data, "father_name_eng", ""),
        "{{BRANCH_SEMESTER}}":      get_safe(data, "branch_semester", ""),
        "{{MOBILE_NO}}":            get_safe(data, "mobile_no", ""),
        "{{INSTITUTE_NAME}}":       get_safe(data, "institute_name", ""),
        "{{GROUP_NAME}}":           get_safe(data, "group_name", ""),
        "{{GROUP_DIRECTOR}}":       get_safe(data, "group_director", ""),
        "{{DIV_HEAD_COORDINATOR}}": get_safe(data, "div_head_coordinator", ""),
        "{{PROJECT_TITLE}}":        get_safe(data, "project_title", ""),
        "{{WORK_DESCRIPTION}}":     get_safe(data, "work_description", ""),
        "{{DATE_OF_JOINING}}":      get_safe(data, "date_of_joining", ""),
        "{{DATE_OF_COMPLETION}}":   get_safe(data, "date_of_completion", ""),
        "{{ATTENDANCE_GRADE}}":     get_safe(data, "attendance_grade", ""),
        "{{PERFORMANCE_RATING}}":   get_safe(data, "performance_rating", ""),
        "{{REPORT_SUBMITTED}}":     get_safe(data, "report_submitted", ""),
        "{{REMARKS}}":              get_safe(data, "remarks", ""),
        "{{RECOMMENDATION}}":       get_safe(data, "recommendation", "")
    }
    
    # Replace in standard paragraphs
    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, placeholders)
    
    # Replace inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para, placeholders)
                    
    # Save the output file
    output_dir = os.path.join(base_dir, "generated_notices")
    os.makedirs(output_dir, exist_ok=True)
    
    # Format a filename
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    clean_name = data.get("student_name_eng", "Trainee").replace(" ", "_")
    filename = f"Appraisal_Proforma_{clean_name}_{timestamp}.docx"
    output_path = os.path.join(output_dir, filename)
    
    if has_unreplaced_placeholders(doc):
        print("WARNING: Unreplaced placeholders detected in generated document!")
    doc.save(output_path)
    return output_path, filename


def generate_coordinator_nomination(master_id, data):
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    # We load the masters list to find the template filename
    masters_file = os.path.join(base_dir, "masters.json")
    with open(masters_file, "r", encoding="utf-8") as f:
        masters = json.load(f)["masters"]
        
    master = next((m for m in masters if m["id"] == master_id), None)
    if not master:
        raise ValueError(f"Master template with ID {master_id} not found!")
        
    template_path = os.path.join(base_dir, "masters", master["filename"])
    doc = Document(template_path)
    
    placeholders = {
        "TRAINING_SESSION": data.get("training_session", ""),
        "REF_NO": data.get("ref_no", ""),
        "DATE": data.get("date", ""),
        "SIGNATORY_NAME": data.get("signatory_name", ""),
        "SIGNATORY_DESIG": data.get("signatory_desig", ""),
        "RECIPIENT_GROUP": data.get("recipient_group", "")
    }
    
    # 1. Replace paragraph placeholders
    for para in doc.paragraphs:
        for key, val in placeholders.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(val))
                
    # 2. Populate Table 0 (Trainees List)
    trainees = data.get("trainees", [])
    
    if len(doc.tables) >= 2:
        table_0 = doc.tables[0]
        # Preserve header (Row 0) and remove all other rows
        while len(table_0.rows) > 1:
            table_0._tbl.remove(table_0.rows[1]._tr)
            
        for idx, t in enumerate(trainees):
            row_cells = table_0.add_row().cells
            row_cells[0].text = f"{idx+1}."
            row_cells[1].text = t.get("name", "")
            row_cells[2].text = t.get("father_name", "")
            row_cells[3].text = t.get("branch", "")
            row_cells[4].text = t.get("group", "")
            
            # Format text in table_0 cells
            for col_idx, cell in enumerate(row_cells):
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 4] else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
        apply_table_formatting(table_0)
                        
        # 3. Populate Table 1 (Coordinator Details)
        table_1 = doc.tables[1]
        # Preserve header (Row 0) and remove all other rows
        while len(table_1.rows) > 1:
            table_1._tbl.remove(table_1.rows[1]._tr)
            
        for idx, t in enumerate(trainees):
            row_cells = table_1.add_row().cells
            row_cells[0].text = f"{idx+1}."
            row_cells[1].text = t.get("name", "")
            row_cells[2].text = "" # Blank for coordinator to fill
            row_cells[3].text = "" # Blank for coordinator to fill
            row_cells[4].text = "" # Blank for coordinator to fill
            
            # Format text in table_1 cells
            for col_idx, cell in enumerate(row_cells):
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
        apply_table_formatting(table_1)
                        
    output_dir = os.path.join(base_dir, "generated_notices")
    os.makedirs(output_dir, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    clean_name = master["name"].replace(" ", "_").replace("/", "_")
    filename = f"{clean_name}_{timestamp}.docx"
    output_path = os.path.join(output_dir, filename)
    if has_unreplaced_placeholders(doc):
        print("WARNING: Unreplaced placeholders detected in generated document!")
    doc.save(output_path)
    return output_path, filename


def generate_internship_noting(master_id, data):
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    # We load the masters list to find the template filename
    masters_file = os.path.join(base_dir, "masters.json")
    with open(masters_file, "r", encoding="utf-8") as f:
        masters = json.load(f)["masters"]
        
    master = next((m for m in masters if m["id"] == master_id), None)
    if not master:
        raise ValueError(f"Master template with ID {master_id} not found!")
        
    template_path = os.path.join(base_dir, "masters", master["filename"])
    doc = Document(template_path)
    
    # 1. Standard variable replacements
    placeholders = {}
    for var in master.get("variables", []):
        placeholders[f"{{{{{var}}}}}"] = data.get(var.lower(), "")
        
    # Replace in standard paragraphs
    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, placeholders)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para, placeholders)
                
    # 2. Populate dynamic tables
    rows_data = data.get("rows", [])
    
    if master_id == "master_008" and len(doc.tables) >= 2:
        # Table 0: S.No, Trainee Name, Branch, Group, Bank Details, Remarks
        table_0 = doc.tables[0]
        while len(table_0.rows) > 1:
            table_0._tbl.remove(table_0.rows[1]._tr)
        for idx, r in enumerate(rows_data):
            row_cells = table_0.add_row().cells
            row_cells[0].text = f"{idx+1}"
            row_cells[1].text = r.get("name", "")
            row_cells[2].text = r.get("branch", "")
            row_cells[3].text = r.get("group", "")
            row_cells[4].text = r.get("bank_details", "")
            row_cells[5].text = r.get("remarks", "")
            
            for c_idx, cell in enumerate(row_cells):
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
        apply_table_formatting(table_0)
                        
        # Table 1: S.No, Trainee Name, Month 1, Month 2, Month 3
        table_1 = doc.tables[1]
        while len(table_1.rows) > 1:
            table_1._tbl.remove(table_1.rows[1]._tr)
        for idx, r in enumerate(rows_data):
            row_cells = table_1.add_row().cells
            row_cells[0].text = f"{idx+1}"
            row_cells[1].text = r.get("name", "")
            row_cells[2].text = "" # blank for coordinator to fill
            row_cells[3].text = ""
            row_cells[4].text = ""
            
            for c_idx, cell in enumerate(row_cells):
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
        apply_table_formatting(table_1)
                        
    elif master_id == "master_009" and len(doc.tables) >= 1:
        # Table 0: S.No, Branch, Requirement, Application Received, Remarks
        table_0 = doc.tables[0]
        while len(table_0.rows) > 1:
            table_0._tbl.remove(table_0.rows[1]._tr)
        for idx, r in enumerate(rows_data):
            row_cells = table_0.add_row().cells
            row_cells[0].text = r.get("sno", f"{idx+1}")
            row_cells[1].text = r.get("branch", "")
            row_cells[2].text = r.get("requirement", "")
            row_cells[3].text = r.get("applications_received", "")
            row_cells[4].text = r.get("remarks", "")
            
            for c_idx, cell in enumerate(row_cells):
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 2, 3] else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
        apply_table_formatting(table_0)
                        
    elif master_id == "master_011" and len(doc.tables) >= 1:
        # Table 0: S.No, Trainee Name, Branch, College, Duration, Remarks
        table_0 = doc.tables[0]
        while len(table_0.rows) > 1:
            table_0._tbl.remove(table_0.rows[1]._tr)
        for idx, r in enumerate(rows_data):
            row_cells = table_0.add_row().cells
            row_cells[0].text = f"{idx+1}"
            row_cells[1].text = r.get("name", "")
            row_cells[2].text = r.get("branch", "")
            row_cells[3].text = r.get("college", "")
            row_cells[4].text = r.get("duration", "")
            row_cells[5].text = r.get("remarks", "")
            
            for c_idx, cell in enumerate(row_cells):
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 4] else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
        apply_table_formatting(table_0)
                        
    # Save the output file
    output_dir = os.path.join(base_dir, "generated_notices")
    os.makedirs(output_dir, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    clean_name = master["name"].replace(" ", "_").replace("/", "_")
    filename = f"{clean_name}_{timestamp}.docx"
    output_path = os.path.join(output_dir, filename)
    if has_unreplaced_placeholders(doc):
        print("WARNING: Unreplaced placeholders detected in generated document!")
    doc.save(output_path)
    return output_path, filename