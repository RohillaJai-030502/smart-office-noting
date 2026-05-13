# ============================================================
# create_master.py — Creates the default master DOCX template
# ============================================================

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top","left","bottom","right","insideH","insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)

def create_default_master():
    os.makedirs("masters", exist_ok=True)
    doc = Document()

    # ── Page Margins ──
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    def add_centered_bold(text, size=12, before=0, after=0, underline=False):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.underline = underline
        set_para_spacing(para, before=before, after=after)
        return para

    # ── HEADER ──
    add_centered_bold("अन्तर कार्यालय नोट / ION", size=12, after=4)
    add_centered_bold("चरम प्राक्षेपिकी अनुसंधान प्रयोगशाला / TBRL", size=11, after=4)
    add_centered_bold("{मानव संसाधन विकास विभाग / HRD}", size=11, after=6)

    # ── SUBJECT ──
    subject_para = doc.add_paragraph()
    subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subject_run = subject_para.add_run(
        "विषय : {{DEGREE}} छात्रों की औद्योगिक प्रशिक्षण"
    )
    subject_run.bold = True
    subject_run.font.size = Pt(11)
    subject_run.underline = True
    set_para_spacing(subject_para, before=2, after=6)

    # ── HINDI BODY ──
    hindi_body = (
        "{{START_MONTH}} {{START_YEAR}} से {{END_MONTH}} {{END_YEAR}} तक प्रशिक्षण प्राप्त "
        "करने के लिये विभिन्न विश्वविद्यालय से {{DEGREE}} के छात्र टी.बी.आर.एल में प्रशिक्षण "
        "प्राप्त करने के लिये आवेदन करेंगे।\n"
        "जो वैज्ञानिक उन्हे प्रशिक्षण देना चाहते हैं, वे अपनी आवश्यकता मानव संसाधन विकास "
        "विभाग को {{LAST_DATE}} तक भेजे । अपनी आवश्यकता को वेAD/TD/GD/JD/PD या वरिष्ठतम "
        "वैज्ञानिक से अग्रेसित अवश्य करायें ।"
    )
    hindi_para = doc.add_paragraph(hindi_body)
    hindi_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    hindi_para.runs[0].font.size = Pt(10)
    set_para_spacing(hindi_para, before=2, after=6)

    # ── ENGLISH BODY 1 ──
    eng1 = (
        "{{DEGREE}} Students from different University interested in getting training at TBRL "
        "during {{START_MONTH}}-{{END_MONTH}} {{END_YEAR}} will submit their applications."
    )
    eng_para1 = doc.add_paragraph(eng1)
    eng_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    eng_para1.runs[0].font.size = Pt(10)
    set_para_spacing(eng_para1, before=2, after=6)

    # ── ENGLISH BODY 2 ──
    eng2 = (
        "Scientists who want to impart training to these students are requested to give their "
        "requirement to HRD Division by {{LAST_DATE}}, forwarded by AD/TD/GD/JD/PD or "
        "senior most scientist of the group."
    )
    eng_para2 = doc.add_paragraph(eng2)
    eng_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    eng_para2.runs[0].font.size = Pt(10)
    set_para_spacing(eng_para2, before=2, after=8)

    # ── SIGNATURE SPACE ──
    for _ in range(3):
        spacer = doc.add_paragraph()
        set_para_spacing(spacer, before=0, after=0, line=8)

    # ── ION + DATE + SIGNATORY TABLE ──
    info_table = doc.add_table(rows=2, cols=2)
    remove_table_borders(info_table)

    left0 = info_table.cell(0, 0).paragraphs[0]
    left0.add_run("ION नं0: {{ION_NUMBER}}").font.size = Pt(10)
    set_para_spacing(left0, before=0, after=2)

    left1 = info_table.cell(1, 0).paragraphs[0]
    left1.add_run("दिनांक:    {{NOTICE_DATE}}").font.size = Pt(10)
    set_para_spacing(left1, before=0, after=2)

    right0 = info_table.cell(0, 1).paragraphs[0]
    right0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_r0 = right0.add_run("({{SIGNATORY_NAME}})")
    run_r0.bold = True
    run_r0.font.size = Pt(10)
    set_para_spacing(right0, before=0, after=2)

    right1 = info_table.cell(1, 1).paragraphs[0]
    right1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right1.add_run("{{SIGNATORY_DESIGNATION}}").font.size = Pt(10)
    set_para_spacing(right1, before=0, after=2)

    # ── TO SECTION ──
    spacer_to = doc.add_paragraph()
    set_para_spacing(spacer_to, before=4, after=2, line=6)

    to_para = doc.add_paragraph()
    to_run = to_para.add_run("To:")
    to_run.bold = True
    to_run.font.size = Pt(10)
    set_para_spacing(to_para, before=2, after=0)

    to_all = doc.add_paragraph("        All TD's / GD's / PD's")
    to_all.runs[0].font.size = Pt(10)
    set_para_spacing(to_all, before=0, after=6)

    # ── DEPARTMENTS TABLE PLACEHOLDER ──
    dept_para = doc.add_paragraph("{{DEPARTMENTS_TABLE}}")
    dept_para.runs[0].font.size = Pt(10)
    set_para_spacing(dept_para, before=2, after=6)

    # ── NOTE ──
    note_para = doc.add_paragraph()
    note_run = note_para.add_run(
        "Note: Performa appended overleaf (the soft copy of the performa "
        "can be downloaded from TBRL portal)"
    )
    note_run.underline = True
    note_run.font.size = Pt(9)
    set_para_spacing(note_para, before=4, after=0)

    # ── SAVE ──
    filepath = os.path.join("masters", "master_001.docx")
    doc.save(filepath)
    print(f"✅ Master template created: {filepath}")
    return filepath

if __name__ == "__main__":
    create_default_master()