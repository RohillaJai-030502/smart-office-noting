from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import math
import shutil
from datetime import datetime
from logic.docx_utils import get_safe, replace_placeholders_in_paragraph, has_unreplaced_placeholders, get_unique_timestamp

OUTPUT_DIR = "generated_notices"

def set_para_spacing(para, before=0, after=0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

def set_table_borders(table):
    """Manually add visible borders to a table using raw XML"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Create the borders element
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')       # Border thickness
        border.set(qn('w:space'), '0')    # Spacing
        border.set(qn('w:color'), '000000') # Black color
        tblBorders.append(border)
        
    tblPr.append(tblBorders)

def build_departments_table(doc, departments):
    """Build a proper Word table for departments"""
    num_cols = min(5, math.ceil(math.sqrt(len(departments) * 2)))
    num_rows = math.ceil(len(departments) / num_cols)

    # Pad departments
    padded = departments + [""] * (num_rows * num_cols - len(departments))

    table = doc.add_table(rows=num_rows, cols=num_cols)
    
    # Use our custom XML border function instead of Word styles
    set_table_borders(table)

    for i, dept in enumerate(padded):
        row = i // num_cols
        col = i % num_cols
        cell = table.cell(row, col)
        cell_para = cell.paragraphs[0]
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell_para.add_run(dept)
        run.font.size = Pt(9)
        set_para_spacing(cell_para, before=2, after=2)

    return table

def generate_ion(master_filename, data):
    """
    Universal Generator: Fills placeholders in master DOCX (Static & Dynamic)
    and saves as new document.
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    master_path = os.path.join("masters", master_filename)
    if not os.path.exists(master_path):
        raise FileNotFoundError(f"Master not found: {master_path}")

    timestamp = get_unique_timestamp()
    
    # Smart Filename Generation: Checks if it's the legacy ION form or a new dynamic form
    if data.get("degree"):
        degree_safe = data['degree'].replace('/', '_')
        filename = f"ION_{degree_safe}_{data.get('start_month', '')}_{data.get('end_year', '')}_{timestamp}.docx"
    else:
        # Generic name for dynamic forms
        filename = f"Generated_Doc_{timestamp}.docx"
        
    output_path = os.path.join(OUTPUT_DIR, filename)
    shutil.copy2(master_path, output_path)

    # Open copied document
    doc = Document(output_path)

    # Build placeholders dict (Preserves legacy mappings for backward compatibility)
    placeholders = {
        "{{DEGREE}}":                   get_safe(data, "degree", ""),
        "{{START_MONTH}}":              get_safe(data, "start_month", ""),
        "{{START_YEAR}}":               get_safe(data, "start_year", ""),
        "{{END_MONTH}}":                get_safe(data, "end_month", ""),
        "{{END_YEAR}}":                 get_safe(data, "end_year", ""),
        "{{LAST_DATE}}":                get_safe(data, "last_date", ""),
        "{{ION_NUMBER}}":               get_safe(data, "ion_number", ""),
        "{{NOTICE_DATE}}":              get_safe(data, "notice_date", ""),
        "{{SIGNATORY_NAME}}":           get_safe(data, "signatory_name", ""),
        "{{SIGNATORY_DESIGNATION}}":    get_safe(data, "signatory_designation", ""),
    }

    # 🚀 DYNAMIC INJECTION: Loops through any new variables sent from the web form
    for key, value in data.items():
        if value is not None:
            # Maps the dictionary key 'employee_name' to the document tag '{{employee_name}}'
            placeholders[f"{{{{{key}}}}}"] = str(value)

    # Replace placeholders in all paragraphs
    dept_para_index = None
    for i, para in enumerate(doc.paragraphs):
        if "{{DEPARTMENTS_TABLE}}" in para.text:
            dept_para_index = i
        else:
            replace_placeholders_in_paragraph(para, placeholders)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para, placeholders)

    # Handle legacy departments table placeholder
    if dept_para_index is not None:
        dept_para = doc.paragraphs[dept_para_index]
        departments = data.get("departments", [])

        if departments:
            # Clear the placeholder text
            for run in dept_para.runs:
                run.text = ""

            # Insert table after the placeholder paragraph
            dept_para._element.addnext(
                build_departments_table(doc, departments)._element
            )

    if has_unreplaced_placeholders(doc):
        print("WARNING: Unreplaced placeholders detected in the generated ION document!")

    doc.save(output_path)
    return output_path, filename