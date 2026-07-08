# ============================================================
# app.py — ION Generator Flask Application (Upgraded)
# ============================================================

from flask import Flask, render_template, request, send_file, redirect, url_for, session, jsonify, flash
from werkzeug.utils import secure_filename
from logic.ion_generator import generate_ion
from logic.create_master import create_default_master
from logic.noting_generator import generate_tbrl_noting, generate_lecture_noting, generate_dgmss_noting, generate_fee_noting, generate_cancellation_noting, generate_date_amendment_fax, generate_mayurpankh_erp_fax, generate_appraisal_proforma, generate_coordinator_nomination, generate_internship_noting  # NEW: TBRL Noting/Fax/ERP Generator Import
import os
import json
import shutil
import re
import docx
from datetime import datetime

import secrets

app = Flask(__name__)

SECRET_KEY_FILE = "secret_key.txt"
if not os.path.exists(SECRET_KEY_FILE):
    with open(SECRET_KEY_FILE, "w") as f:
        f.write(secrets.token_hex(32))
with open(SECRET_KEY_FILE, "r") as f:
    app.secret_key = f.read().strip()

# ── File Paths ──
DATA_FILE     = "data.json"
HISTORY_FILE  = "history.json"
DEFAULTS_FILE = "defaults.json"
MASTERS_FILE  = "masters.json"
CONFIG_FILE   = "config.json"

DEGREE_OPTIONS = ["B.Tech/B.E.", "M.Tech", "B.Sc", "M.Sc", "PhD"]

MONTHS = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December"
]

# ══════════════════════════════════════
# HELPERS
# ══════════════════════════════════════
def load_json(filepath, default):
    if os.path.exists(filepath):
        with open(filepath, "r") as f:
            return json.load(f)
    return default

def save_json(filepath, data):
    import tempfile
    dir_name = os.path.dirname(os.path.abspath(filepath)) or "."
    fd, tmp_path = tempfile.mkstemp(dir=dir_name)
    try:
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        os.replace(tmp_path, filepath)
    except Exception:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        raise

def load_departments():
    return load_json(DATA_FILE, {"departments": []})["departments"]

def save_departments(departments):
    save_json(DATA_FILE, {"departments": departments})

def load_history():
    return load_json(HISTORY_FILE, {"history": []})["history"]

def save_history(entry):
    history = load_history()
    history.insert(0, entry)
    history = history[:20]
    save_json(HISTORY_FILE, {"history": history})

def load_defaults():
    return load_json(DEFAULTS_FILE, {})

def load_courses():
    try:
        with open("courses.json", "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []

def load_columns_map():
    try:
        with open("columns.json", "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def load_masters():
    masters = load_json(MASTERS_FILE, {"masters": []})["masters"]
    valid_masters = []
    for m in masters:
        filepath = os.path.join("masters", m["filename"])
        if os.path.exists(filepath):
            valid_masters.append(m)
    return valid_masters

def save_masters(masters):
    save_json(MASTERS_FILE, {"masters": masters})

def load_config():
    config = load_json(CONFIG_FILE, {"password": "Tbrl0000"})
    if "password_hash" not in config:
        from werkzeug.security import generate_password_hash
        old_password = config.get("password", "Tbrl0000")
        config["password_hash"] = generate_password_hash(old_password)
        if "password" in config:
            del config["password"]
        save_json(CONFIG_FILE, config)
    return config

from werkzeug.security import generate_password_hash, check_password_hash

def check_password(password):
    config = load_config()
    return check_password_hash(config.get("password_hash", ""), password)

def form_get(field, default=""):
    """request.form.get() that respects the default for blank/empty values too."""
    val = request.form.get(field)
    return default if val is None or val == "" else val

def safe_int(value, default=1):
    """Safely converts input to integer, falling back to default on error."""
    try:
        return int(value)
    except (TypeError, ValueError):
        return default

# Initialize config and auto-hash password on server startup
load_config()

def parse_date_safe(value, fmt="%Y-%m-%d"):
    """Parses a date string safely, raising a ValueError on formatting error."""
    if not value:
        raise ValueError("Missing or empty date value.")
    try:
        return datetime.strptime(str(value).strip(), fmt)
    except (TypeError, ValueError) as e:
        raise ValueError(f"Invalid date format: '{value}'. Expected '{fmt}'.") from e

@app.errorhandler(Exception)
def handle_any_error(e):
    app.logger.exception("Unhandled error encountered")
    return render_template("error.html", message="Something went wrong generating your document. Please check your form entries and try again."), 500

# ── NEW HELPER: Auto-Extract Variables ──
def extract_variables_from_docx(filepath):
    """Reads a .docx file and extracts all variables inside {{ }}"""
    doc = docx.Document(filepath)
    variables = set()
    pattern = r'\{\{(.+?)\}\}'
    
    for para in doc.paragraphs:
        matches = re.findall(pattern, para.text)
        for match in matches:
            variables.add(match.strip())
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(pattern, cell.text)
                for match in matches:
                    variables.add(match.strip())
                    
    return list(variables)

def get_master_safe(masters, master_id):
    """Safely retrieves a master template by ID, falling back to the first one or None."""
    if not masters:
        return None
    master = next((m for m in masters if m["id"] == master_id), None)
    return master if master is not None else masters[0]

# ══════════════════════════════════════
# DASHBOARD
# ══════════════════════════════════════
@app.route("/")
def index():
    return render_template("dashboard.html",
        history=load_history()[:5],
        masters=load_masters()
    )

# ══════════════════════════════════════
# ION NOTICE FORM (Legacy / Original Form)
# ══════════════════════════════════════
@app.route("/ion-notice", methods=["GET"])
def ion_notice():
    saved    = session.pop("form_data", {})
    defaults = load_defaults()
    masters  = load_masters()
    
    # Determine which master is selected (first default fillable master, or first fillable master)
    fillable_masters = [m for m in masters if m.get("form_type") == "fillable"]
    selected_master_id = None
    if saved and saved.get("master_id"):
        selected_master_id = saved["master_id"]
    else:
        # Find first default fillable master
        default_master = next((m for m in fillable_masters if m.get("is_default")), None)
        if default_master:
            selected_master_id = default_master["id"]
        elif fillable_masters:
            selected_master_id = fillable_masters[0]["id"]
            
    return render_template(
        "ion_form.html",
        departments=load_departments(),
        degree_options=DEGREE_OPTIONS,
        months=MONTHS,
        saved=saved,
        defaults=defaults,
        masters=masters,
        selected_master_id=selected_master_id
    )

@app.route("/save-defaults", methods=["POST"])
def save_defaults_route():
    defaults = {
        "signatory_name":        form_get("signatory_name", ""),
        "signatory_designation": form_get("signatory_designation", ""),
    }
    save_json(DEFAULTS_FILE, defaults)
    session["form_data"] = {
        "master_id":             form_get("master_id"),
        "degree":                form_get("degree"),
        "start_month":           form_get("start_month"),
        "start_year":            form_get("start_year"),
        "end_month":             form_get("end_month"),
        "end_year":              form_get("end_year"),
        "last_date":             form_get("last_date"),
        "ion_number":            form_get("ion_number"),
        "notice_date":           form_get("notice_date"),
        "signatory_name":        form_get("signatory_name"),
        "signatory_designation": form_get("signatory_designation"),
        "departments":           request.form.getlist("departments"),
    }
    return redirect(url_for("ion_notice"))

@app.route("/preview", methods=["POST"])
def preview():
    departments = request.form.getlist("departments")
    num_cols    = 5
    num_rows    = -(-len(departments) // num_cols)
    padded      = departments + [""] * (num_rows * num_cols - len(departments))
    dept_rows   = [padded[r * num_cols:(r + 1) * num_cols] for r in range(num_rows)]

    data = {
        "master_id":             form_get("master_id"),
        "degree":                form_get("degree"),
        "start_month":           form_get("start_month"),
        "start_year":            form_get("start_year"),
        "end_month":             form_get("end_month"),
        "end_year":              form_get("end_year"),
        "last_date":             form_get("last_date"),
        "ion_number":            form_get("ion_number"),
        "notice_date":           form_get("notice_date"),
        "signatory_name":        form_get("signatory_name"),
        "signatory_designation": form_get("signatory_designation"),
        "departments":           departments,
        "dept_rows":             dept_rows,
    }
    session["form_data"] = data
    return render_template("ion_preview.html", data=data)

@app.route("/edit", methods=["POST"])
def edit():
    session["form_data"] = {
        "master_id":             form_get("master_id"),
        "degree":                form_get("degree"),
        "start_month":           form_get("start_month"),
        "start_year":            form_get("start_year"),
        "end_month":             form_get("end_month"),
        "end_year":              form_get("end_year"),
        "last_date":             form_get("last_date"),
        "ion_number":            form_get("ion_number"),
        "notice_date":           form_get("notice_date"),
        "signatory_name":        form_get("signatory_name"),
        "signatory_designation": form_get("signatory_designation"),
        "departments":           request.form.getlist("departments"),
    }
    return redirect(url_for("ion_notice"))

@app.route("/generate", methods=["POST"])
def generate():
    masters   = load_masters()
    master_id = form_get("master_id", "master_001")
    master = get_master_safe(masters, master_id)
    if not master:
        flash("No master templates exist yet. Please add one first.")
        return redirect(url_for("masters_manager"))

    data = {
        "degree":                form_get("degree"),
        "start_month":           form_get("start_month"),
        "start_year":            form_get("start_year"),
        "end_month":             form_get("end_month"),
        "end_year":              form_get("end_year"),
        "last_date":             form_get("last_date"),
        "ion_number":            form_get("ion_number"),
        "notice_date":           form_get("notice_date"),
        "signatory_name":        form_get("signatory_name"),
        "signatory_designation": form_get("signatory_designation"),
        "departments":           request.form.getlist("departments"),
    }

    filepath, filename = generate_ion(master["filename"], data)
    save_history({
        "filename":          filename,
        "degree":            data["degree"],
        "period":            f"{data['start_month']} {data['start_year']} - {data['end_month']} {data['end_year']}",
        "generated_at":      datetime.now().strftime("%d %b %Y, %I:%M %p"),
        "departments_count": len(data["departments"]),
        "master_used":       master["name"],
    })
    return send_file(filepath, as_attachment=True)

@app.route("/download", methods=["POST"])
def download():
    masters   = load_masters()
    master_id = form_get("master_id", "master_001")
    master = get_master_safe(masters, master_id)
    if not master:
        flash("No master templates exist yet. Please add one first.")
        return redirect(url_for("masters_manager"))

    data = {
        "degree":                form_get("degree"),
        "start_month":           form_get("start_month"),
        "start_year":            form_get("start_year"),
        "end_month":             form_get("end_month"),
        "end_year":              form_get("end_year"),
        "last_date":             form_get("last_date"),
        "ion_number":            form_get("ion_number"),
        "notice_date":           form_get("notice_date"),
        "signatory_name":        form_get("signatory_name"),
        "signatory_designation": form_get("signatory_designation"),
        "departments":           request.form.getlist("departments"),
    }

    filepath, filename = generate_ion(master["filename"], data)
    save_history({
        "filename":          filename,
        "degree":            data["degree"],
        "period":            f"{data['start_month']} {data['start_year']} - {data['end_month']} {data['end_year']}",
        "generated_at":      datetime.now().strftime("%d %b %Y, %I:%M %p"),
        "departments_count": len(data["departments"]),
        "master_used":       master["name"],
    })
    return send_file(filepath, as_attachment=True)

@app.route("/history/download/<filename>")
def download_history(filename):
    filepath = os.path.join("generated_notices", filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    return "File not found", 404

@app.route("/history/clear", methods=["POST"])
def clear_history():
    save_json(HISTORY_FILE, {"history": []})
    return redirect(url_for("index"))

@app.route("/download-static/<master_id>", methods=["GET"])
def download_static(master_id):
    masters = load_masters()
    master = next((m for m in masters if m["id"] == master_id), None)
    
    if master and master.get("form_type") == "static":
        filepath = os.path.join("masters", master["filename"])
        if os.path.exists(filepath):
            save_history({
                "filename": master["filename"],
                "degree": "Blank Form",
                "period": "N/A",
                "generated_at": datetime.now().strftime("%d %b %Y, %I:%M %p"),
                "departments_count": 0,
                "master_used": master["name"],
            })
            return send_file(filepath, as_attachment=True, download_name=master["filename"])
            
    return "Form not found or is not a static document.", 404

# ══════════════════════════════════════
# DEPARTMENT MANAGER
# ══════════════════════════════════════
@app.route("/departments/add", methods=["POST"])
def add_department():
    name = form_get("new_dept", "").strip().upper()
    if name:
        departments = load_departments()
        if name not in departments:
            departments.append(name)
            save_departments(departments)
    return redirect(url_for("ion_notice"))

@app.route("/departments/delete/<name>", methods=["POST"])
def delete_department(name):
    departments = load_departments()
    departments = [d for d in departments if d != name]
    save_departments(departments)
    return redirect(url_for("ion_notice"))

@app.route("/departments/edit", methods=["POST"])
def edit_department():
    old_name = form_get("old_name", "").strip()
    new_name = form_get("new_name", "").strip().upper()
    if old_name and new_name:
        departments = load_departments()
        departments = [new_name if d == old_name else d for d in departments]
        save_departments(departments)
    return redirect(url_for("ion_notice"))

# ══════════════════════════════════════
# MASTER DOCUMENT MANAGER (Password Protected)
# ══════════════════════════════════════
@app.route("/masters", methods=["GET", "POST"])
def masters_page():
    error = None
    if request.method == "POST":
        password = form_get("password", "")
        if check_password(password):
            session["master_unlocked"] = True
            return redirect(url_for("masters_manager"))
        else:
            error = "❌ Incorrect password. Please try again."
    return render_template("master_password.html", error=error)

@app.route("/masters/manager")
def masters_manager():
    if not session.get("master_unlocked"):
        return redirect(url_for("masters_page"))
    return render_template("master_manager.html", masters=load_masters())

@app.route("/masters/lock")
def masters_lock():
    session.pop("master_unlocked", None)
    return redirect(url_for("index"))

@app.route("/masters/create", methods=["POST"])
def create_master():
    if not session.get("master_unlocked"):
        return redirect(url_for("masters_page"))

    source_id   = form_get("source_id")
    name        = form_get("name", "").strip()
    description = form_get("description", "").strip()

    if not name:
        return redirect(url_for("masters_manager"))

    masters     = load_masters()
    source      = get_master_safe(masters, source_id)
    if not source:
        flash("No master templates exist to copy from.")
        return redirect(url_for("masters_manager"))

    new_id       = f"master_{str(len(masters)+1).zfill(3)}"
    new_filename = f"{new_id}.docx"

    src_path = os.path.join("masters", source["filename"])
    dst_path = os.path.join("masters", new_filename)
    shutil.copy2(src_path, dst_path)

    masters.append({
        "id":          new_id,
        "name":        name,
        "description": description,
        "filename":    new_filename,
        "created_at":  datetime.now().strftime("%Y-%m-%d"),
        "is_default":  False,
        "form_type":   "static" # Ensures backward compatibility 
    })
    save_masters(masters)
    return redirect(url_for("masters_manager"))

# ── NEW ROUTE: Upload Brand New Document & Extract Variables ──
@app.route("/masters/upload-new", methods=["POST"])
def upload_new_master():
    if not session.get("master_unlocked"):
        return redirect(url_for("masters_page"))
    
    if 'file' not in request.files:
        flash("No file provided.")
        return redirect(url_for("masters_manager"))
        
    file = request.files['file']
    name = form_get("name", "Unnamed Master").strip()
    description = form_get("description", "Uploaded via Document Manager").strip()
    
    if not file.filename.lower().endswith('.docx'):
        flash("Only .docx files are accepted.")
        return redirect(url_for("masters_manager"))
        
    masters = load_masters()
    new_id = f"master_{str(len(masters)+1).zfill(3)}"
    safe_filename = secure_filename(f"{new_id}_{file.filename}")
    
    # Save file to masters folder
    os.makedirs("masters", exist_ok=True)
    filepath = os.path.join("masters", safe_filename)
    file.save(filepath)
    
    try:
        # Trigger Variable Extraction
        variables = extract_variables_from_docx(filepath)
    except Exception:
        if os.path.exists(filepath):
            os.remove(filepath)
        flash("The uploaded file is not a valid Word (.docx) document. Please check it and try again.")
        return redirect(url_for("masters_manager"))
        
    # 🧠 SMART CATEGORIZATION LOGIC
    if len(variables) > 0:
        form_type = "dynamic"
        flash_msg = f"✨ Dynamic Master added! Extracted {len(variables)} form fields."
    else:
        form_type = "static"
        flash_msg = "🖨️ Static Template added! No variables found, saving as a printable blank document."
        
        # Append to masters.json
        masters.append({
            "id": new_id,
            "name": name,
            "description": description,
            "filename": safe_filename,
            "created_at": datetime.now().strftime("%Y-%m-%d"),
            "is_default": False,
            "form_type": form_type,  # Now correctly assigns dynamic vs static
            "variables": variables
        })
        save_masters(masters)
        flash(flash_msg)
        
    return redirect(url_for("masters_manager"))

@app.route("/masters/toggle-default/<master_id>", methods=["POST"])
def toggle_default_master(master_id):
    if not session.get("master_unlocked"):
        return redirect(url_for("masters_page"))
    masters = load_masters()
    for m in masters:
        if m["id"] == master_id:
            m["is_default"] = not m.get("is_default", False)
            break
    save_masters(masters)
    return redirect(url_for("masters_manager"))

@app.route("/masters/delete/<master_id>", methods=["POST"])
def delete_master(master_id):
    if not session.get("master_unlocked"):
        return redirect(url_for("masters_page"))
        
    # 🔒 NEW: Secondary password check for deletion
    confirm_password = form_get("confirm_password", "")
    if not check_password(confirm_password):
        flash("❌ Incorrect password! Deletion cancelled.")
        return redirect(url_for("masters_manager"))

    masters = load_masters()
    master  = next((m for m in masters if m["id"] == master_id), None)
    
    if master and not master.get("is_default"):
        filepath = os.path.join("masters", master["filename"])
        if os.path.exists(filepath):
            try:
                os.remove(filepath)
            except OSError:
                pass # Ignore if file was already removed manually
                
        masters = [m for m in masters if m["id"] != master_id]
        save_masters(masters)
        flash(f"🗑️ Master template '{master['name']}' was permanently deleted.")
        
    return redirect(url_for("masters_manager"))

@app.route("/masters/download/<master_id>")
def download_master(master_id):
    if not session.get("master_unlocked"):
        return redirect(url_for("masters_page"))
    masters = load_masters()
    master  = next((m for m in masters if m["id"] == master_id), None)
    if master:
        filepath = os.path.join("masters", master["filename"])
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True,
                           download_name=f"{master['name']}.docx")
    return "Master not found", 404

@app.route("/masters/upload/<master_id>", methods=["POST"])
def upload_master(master_id):
    if not session.get("master_unlocked"):
        return redirect(url_for("masters_page"))
    masters = load_masters()
    master  = next((m for m in masters if m["id"] == master_id), None)
    if master and "file" in request.files:
        file = request.files["file"]
        if not file.filename.lower().endswith(".docx"):
            flash("Only .docx files are accepted.")
            return redirect(url_for("masters_manager"))
            
        filepath = os.path.join("masters", master["filename"])
        temp_filepath = filepath + ".tmp"
        file.save(temp_filepath)
        
        try:
            # Validate docx by parsing it
            extract_variables_from_docx(temp_filepath)
            
            # Replace existing file
            if os.path.exists(filepath):
                os.remove(filepath)
            os.rename(temp_filepath, filepath)
            
            # Re-extract and update variables list in masters.json
            variables = extract_variables_from_docx(filepath)
            master["variables"] = variables
            save_masters(masters)
            flash("Template updated successfully!")
        except Exception:
            if os.path.exists(temp_filepath):
                os.remove(temp_filepath)
            flash("The uploaded file is not a valid Word (.docx) document. Please check it and try again.")
            
    return redirect(url_for("masters_manager"))

# ══════════════════════════════════════
# DYNAMIC FORM RENDERING & SUBMISSION
# ══════════════════════════════════════
@app.route('/dynamic-form/<master_id>')
def dynamic_form(master_id):
    masters = load_masters()
    master = next((m for m in masters if m["id"] == master_id), None)
    
    if not master or master.get("form_type") != "dynamic":
        return "Dynamic form not found or invalid type.", 404
        
    return render_template('dynamic_form.html', template=master)

@app.route('/submit-dynamic', methods=['POST'])
def submit_dynamic():
    master_id = form_get("template_id")
    masters = load_masters()
    master = next((m for m in masters if m["id"] == master_id), None)
    
    if not master:
        return "Master not found", 404
        
    # 1. Grab all the user's answers from the form
    data = {}
    for var in master.get("variables", []):
        data[var] = form_get(var, "")
        
    # 2. Open the Master Document
    master_path = os.path.join("masters", master["filename"])
    doc = docx.Document(master_path)
    
    def replace_text(paragraph, key, value):
        placeholder = f"{{{{{key}}}}}" # Looks for {{variable_name}}
        if placeholder in paragraph.text:
            replaced = False
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
                    replaced = True
            if not replaced:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # 4. Search and Replace in standard text
    for para in doc.paragraphs:
        for k, v in data.items():
            replace_text(para, k, v)
            
    # 5. Search and Replace inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for k, v in data.items():
                        replace_text(para, k, v)
                        
    # 6. Save the new generated file
    os.makedirs("generated_notices", exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    clean_name = master['name'].replace(' ', '_').replace('/', '_').replace('\\', '_')
    filename = f"Dynamic_{clean_name}_{timestamp}.docx"
    output_path = os.path.join("generated_notices", filename)
    doc.save(output_path)
    
    # 7. Add to history
    save_history({
        "filename": filename,
        "degree": "Dynamic Form",
        "period": "N/A",
        "generated_at": datetime.now().strftime("%d %b %Y, %I:%M %p"),
        "departments_count": len(data),
        "master_used": master["name"],
    })
    
    # 8. Send to user
    return send_file(output_path, as_attachment=True)

# ══════════════════════════════════════
# TBRL NOTING GENERATOR (NEW)
# ══════════════════════════════════════
NOTING_CONFIGS = {
    "tbrl": {
        "generator": generate_tbrl_noting,
        "label": "TBRL Noting",
        "master_used": "TBRL Noting Master"
    },
    "lecture": {
        "generator": generate_lecture_noting,
        "label": "Lecture Noting",
        "master_used": "Lecture Noting Master"
    },
    "dgmss": {
        "generator": generate_dgmss_noting,
        "label": "DGMSS Noting",
        "master_used": "DGMSS Noting Master"
    },
    "fee": {
        "generator": generate_fee_noting,
        "label": "Fee Related Noting",
        "master_used": "Fee Related Noting Master"
    },
    "cancellation": {
        "generator": generate_cancellation_noting,
        "label": "Nomination Cancellation",
        "master_used": "Nomination Cancellation Master"
    }
}

def handle_noting_generation(config_key, extra_fields_fn=None):
    cfg = NOTING_CONFIGS[config_key]

    # 1. Update Signatory Defaults in Memory
    defaults = load_defaults()
    defaults["sig1_name"] = form_get("sig1_name", "")
    defaults["sig1_desig"] = form_get("sig1_desig", "")
    defaults["sig2_name"] = form_get("sig2_name", "")
    defaults["sig2_desig"] = form_get("sig2_desig", "")
    save_json(DEFAULTS_FILE, defaults)

    # 2. Extract Table Configuration & Data
    columns = request.form.getlist("table_columns")
    num_rows = safe_int(form_get("num_rows", 1), 1)
    
    nominees = []
    for i in range(num_rows):
        row_data = {}
        for col in columns:
            row_data[col] = form_get(f"{col}_{i}", "")
        nominees.append(row_data)

    # 3. Handle 'Other' Course Type
    course_type = form_get("course_type")
    if course_type == "Others":
        course_type = form_get("custom_course_type", "Course")

    ref_no = form_get("form_no") or form_get("ref_no")
    lab_name = form_get("lab_name", "टीबीआरएल")
    
    references_raw = form_get("references_json")
    references = []
    if references_raw:
        try:
            references = json.loads(references_raw)
        except Exception:
            pass
    if not references:
        references = [{
            "source": ref_no or "",
            "date": form_get("ref_date", "")
        }]

    # 4. Compile all data
    data = {
        "ref_no": ref_no,
        "subject_hindi": form_get("subject_hindi"),
        "subject_english": form_get("subject_english"),
        "reference_text": form_get("reference_text"),
        "ref_date": form_get("ref_date"),
        "start_date": form_get("start_date"),
        "end_date": form_get("end_date"),
        "org_institute": form_get("org_institute"),
        "course_title": form_get("course_title"),
        "course_type": course_type,
        "group_name": form_get("group_name"),
        "sig1_name": defaults["sig1_name"],
        "sig1_desig": defaults["sig1_desig"],
        "sig2_name": defaults["sig2_name"],
        "sig2_desig": defaults["sig2_desig"],
        "columns": columns,
        "nominees": nominees,
        "lab_name": lab_name,
        "references": references
    }
    
    if extra_fields_fn:
        data.update(extra_fields_fn(request.form))

    # 5. Generate Document
    filepath, filename = cfg["generator"](data)
    
    # 6. Save to Dashboard History
    save_history({
        "filename": filename,
        "degree": cfg["label"],
        "period": f"{data.get('start_date', '')} to {data.get('end_date', '')}",
        "generated_at": datetime.now().strftime("%d %b %Y, %I:%M %p"),
        "departments_count": len(nominees),
        "master_used": cfg["master_used"]
    })
    
    return send_file(filepath, as_attachment=True)

@app.route("/tbrl-noting", methods=["GET"])
def tbrl_noting():
    defaults = load_defaults()
    groups = ["AFTD", "ADS", "BIDS", "BEHI", "SS", "TELIC", "PPG", "QMG", "ETF", "PC", "WHD", "EXPD", "WHT&E", "ARISE", "PCD", "AIG", "RTRS", "S&D", "R&QA", "WKS", "SEED", "CERBERUS", "DPB", "HSP", "I2G", "HRDD"]
    courses = load_courses()
    default_columns = load_columns_map().get("tbrl_noting", [])
    return render_template("tbrl_noting_form.html", defaults=defaults, groups=groups, courses=courses, default_columns=default_columns)

@app.route("/generate-noting", methods=["POST"])
def generate_noting():
    return handle_noting_generation("tbrl")

# ══════════════════════════════════════
# LECTURE NOTING GENERATOR
# ══════════════════════════════════════
@app.route("/lecture-noting", methods=["GET"])
def lecture_noting():
    defaults = load_defaults()
    groups = ["AFTD", "ADS", "BIDS", "BEHI", "SS", "TELIC", "PPG", "QMG", "ETF", "PC", "WHD", "EXPD", "WHT&E", "ARISE", "PCD", "AIG", "RTRS", "S&D", "R&QA", "WKS", "SEED", "CERBERUS", "DPB", "HSP", "I2G", "HRDD"]
    courses = load_courses()
    default_columns = load_columns_map().get("lecture_noting", [])
    return render_template("lecture_noting_form.html", defaults=defaults, groups=groups, courses=courses, default_columns=default_columns)


@app.route("/generate-lecture-noting", methods=["POST"])
def generate_lecture_noting_route():
    return handle_noting_generation("lecture")

# ══════════════════════════════════════
# DGMSS NOTING GENERATOR
# ══════════════════════════════════════
@app.route("/dgmss-noting", methods=["GET"])
def dgmss_noting():
    defaults = load_defaults()
    groups = ["AFTD", "ADS", "BIDS", "BEHI", "SS", "TELIC", "PPG", "QMG", "ETF", "PC", "WHD", "EXPD", "WHT&E", "ARISE", "PCD", "AIG", "RTRS", "S&D", "R&QA", "WKS", "SEED", "CERBERUS", "DPB", "HSP", "I2G", "HRDD", "BTS"]
    courses = load_courses()
    default_columns = load_columns_map().get("dgmss_noting", [])
    return render_template("dgmss_noting_form.html", defaults=defaults, groups=groups, courses=courses, default_columns=default_columns)


@app.route("/generate-dgmss-noting", methods=["POST"])
def generate_dgmss_noting_route():
    return handle_noting_generation("dgmss")

# ══════════════════════════════════════
@app.route("/fee-noting", methods=["GET"])
def fee_noting():
    defaults = load_defaults()
    groups = ["AFTD", "ADS", "BIDS", "BEHI", "SS", "TELIC", "PPG", "QMG", "ETF", "PC", "WHD", "EXPD", "WHT&E", "ARISE", "PCD", "AIG", "RTRS", "S&D", "R&QA", "WKS", "SEED", "CERBERUS", "DPB", "HSP", "I2G", "HRDD", "BTS"]
    courses = load_courses()
    default_columns = load_columns_map().get("fee_noting", [])
    return render_template("fee_noting_form.html", defaults=defaults, groups=groups, courses=courses, default_columns=default_columns)


@app.route("/generate-fee-noting", methods=["POST"])
def generate_fee_noting_route():
    def fee_extra(form):
        return {"ref_mail_date": form_get("ref_mail_date")}
    return handle_noting_generation("fee", fee_extra)

# ══════════════════════════════════════
@app.route("/cancellation-noting", methods=["GET"])
def cancellation_noting():
    defaults = load_defaults()
    groups = ["AFTD", "ADS", "BIDS", "BEHI", "SS", "TELIC", "PPG", "QMG", "ETF", "PC", "WHD", "EXPD", "WHT&E", "ARISE", "PCD", "AIG", "RTRS", "S&D", "R&QA", "WKS", "SEED", "CERBERUS", "DPB", "HSP", "I2G", "HRDD", "BTS"]
    courses = load_courses()
    default_columns = load_columns_map().get("cancellation_noting", [])
    return render_template("cancellation_noting_form.html", defaults=defaults, groups=groups, courses=courses, default_columns=default_columns)


@app.route("/generate-cancellation-noting", methods=["POST"])
def generate_cancellation_noting_route():
    def cancellation_extra(form):
        return {
            "ref_mail_date": form_get("ref_mail_date"),
            "ion_ref_source": form_get("ion_ref_source"),
            "ion_ref_date": form_get("ion_ref_date"),
            "cancel_nominee_name": form_get("cancel_nominee_name"),
            "cancel_group_name": form_get("cancel_group_name"),
            "cancel_reason": form_get("cancel_reason"),
            "course_type_short": form_get("course_type_short"),
        }
    return handle_noting_generation("cancellation", cancellation_extra)

# ═════════════════════════════
@app.route("/date-amendment-fax", methods=["GET"])
def date_amendment_fax():
    defaults = load_defaults()
    courses = load_courses()
    default_columns = load_columns_map().get("date_amendment_fax", [])
    return render_template("date_amendment_fax_form.html", defaults=defaults, courses=courses, default_columns=default_columns)


@app.route("/generate-date-amendment-fax", methods=["POST"])
def generate_date_amendment_fax_route():
    # 1. Extract Table Configuration & Data
    for_columns = request.form.getlist("for_columns")
    read_columns = request.form.getlist("read_columns")
    for_rows_count = safe_int(form_get("for_rows_count", 1), 1)
    read_rows_count = safe_int(form_get("read_rows_count", 1), 1)

    for_rows = []
    for i in range(for_rows_count):
        row_data = {}
        for col in for_columns:
            row_data[col] = form_get(f"for_{col}_{i}", "")
        for_rows.append(row_data)

    read_rows = []
    for i in range(read_rows_count):
        row_data = {}
        for col in read_columns:
            row_data[col] = form_get(f"read_{col}_{i}", "")
        read_rows.append(row_data)

    # 2. Assemble document payload
    data = {
        "ref_no": form_get("ref_no", ""),
        "fax_no": form_get("fax_no", ""),
        "ref_date": form_get("ref_date", ""),
        "to_text": form_get("to_text", ""),
        "subject_hindi": form_get("subject_hindi", ""),
        "subject_english": form_get("subject_english", ""),
        "ref_text": form_get("ref_text", ""),
        "ref_to_hindi": form_get("ref_to_hindi", ""),
        "ref_to_eng": form_get("ref_to_eng", ""),
        "num_courses_hindi": form_get("num_courses_hindi", "तीन"),
        "num_courses_eng": form_get("num_courses_eng", "three"),
        "courses_plural_eng": form_get("courses_plural_eng", "Courses"),
        "amended_courses_hindi": form_get("amended_courses_hindi", "एक"),
        "amended_courses_eng": form_get("amended_courses_eng", "one"),
        "for_columns": for_columns,
        "for_rows": for_rows,
        "read_columns": read_columns,
        "read_rows": read_rows,
        "sig_name": form_get("sig_name", ""),
        "sig_desig": form_get("sig_desig", ""),
        "for_director": form_get("for_director", "")
    }

    # Generate Fax document
    filepath, filename = generate_date_amendment_fax(data)

    # Save to dashboard history
    save_history({
        "filename": filename,
        "degree": f"Date Amendment Fax ({form_get('course_type_select', 'C.E.P.')})",
        "period": "N/A",
        "generated_at": datetime.now().strftime("%d %b %Y, %I:%M %p"),
        "departments_count": for_rows_count,
        "master_used": "Date Amendment Fax Master"
    })

    return send_file(filepath, as_attachment=True)


# MAYURPANKH ERP FAX GENERATOR
# ═════════════════════════════
@app.route("/mayurpankh-erp", methods=["GET"])
def mayurpankh_erp():
    defaults = load_defaults()
    courses = load_courses()
    default_columns = load_columns_map().get("mayurpankh_erp", [])
    return render_template("mayurpankh_erp_form.html", defaults=defaults, courses=courses, default_columns=default_columns)


@app.route("/generate-mayurpankh-erp-fax", methods=["POST"])
def generate_mayurpankh_erp_fax_route():
    # 1. Extract Table Configuration & Data
    columns = request.form.getlist("columns")
    rows_count = safe_int(form_get("rows_count", 1), 1)

    rows = []
    for i in range(rows_count):
        row_data = {}
        for col in columns:
            row_data[col] = form_get(f"{col}_{i}", "")
        rows.append(row_data)

    # 2. Assemble document payload
    data = {
        "ref_no": form_get("ref_no", ""),
        "fax_no": form_get("fax_no", ""),
        "ref_date": form_get("ref_date", ""),
        "to_text": form_get("to_text", ""),
        "subject_hindi": form_get("subject_hindi", ""),
        "subject_english": form_get("subject_english", ""),
        "ref_text": form_get("ref_text", ""),
        "attn_text": form_get("attn_text", ""),
        "body_hindi": form_get("body_hindi", ""),
        "body_english": form_get("body_english", ""),
        "confirm_text": form_get("confirm_text", ""),
        "columns": columns,
        "rows": rows,
        "sig_name": form_get("sig_name", ""),
        "sig_desig": form_get("sig_desig", ""),
        "for_director": form_get("for_director", "")
    }

    # Generate Fax document
    filepath, filename = generate_mayurpankh_erp_fax(data)

    # Save to dashboard history
    save_history({
        "filename": filename,
        "degree": "Mayurpankh ERP Fax",
        "period": "N/A",
        "generated_at": datetime.now().strftime("%d %b %Y, %I:%M %p"),
        "departments_count": rows_count,
        "master_used": "Mayurpankh ERP Fax Master"
    })

    return send_file(filepath, as_attachment=True)


# APPRAISAL PROFORMA GENERATOR
# ══════════════════════════════════════
@app.route("/appraisal-proforma", methods=["GET"])
def appraisal_proforma():
    defaults = load_defaults()
    groups = ["AFTD", "ADS", "BIDS", "BEHI", "SS", "TELIC", "PPG", "QMG", "ETF", "PC", "WHD", "EXPD", "WHT&E", "ARISE", "PCD", "AIG", "RTRS", "S&D", "R&QA", "WKS", "SEED", "CERBERUS", "DPB", "HSP", "I2G", "HRDD", "BTS"]
    return render_template("appraisal_form.html", defaults=defaults, groups=groups)


@app.route("/generate-appraisal-proforma", methods=["POST"])
def generate_appraisal_proforma_route():
    data = {
        "student_name_hindi":   form_get("student_name_hindi", ""),
        "father_name_hindi":    form_get("father_name_hindi", ""),
        "student_name_eng":     form_get("student_name_eng", ""),
        "father_name_eng":      form_get("father_name_eng", ""),
        "branch_semester":      form_get("branch_semester", ""),
        "mobile_no":            form_get("mobile_no", ""),
        "institute_name":       form_get("institute_name", ""),
        "group_name":           form_get("group_name", ""),
        "group_director":       form_get("group_director", ""),
        "div_head_coordinator": form_get("div_head_coordinator", ""),
        "project_title":        form_get("project_title", ""),
        "work_description":     form_get("work_description", ""),
        "date_of_joining":      form_get("date_of_joining", ""),
        "date_of_completion":   form_get("date_of_completion", ""),
        "attendance_grade":     form_get("attendance_grade", ""),
        "performance_rating":   form_get("performance_rating", ""),
        "report_submitted":     form_get("report_submitted", ""),
        "remarks":              form_get("remarks", ""),
        "recommendation":       form_get("recommendation", "")
    }

    filepath, filename = generate_appraisal_proforma(data)

    save_history({
        "filename":          filename,
        "degree":            data["branch_semester"],
        "period":            f"{data['date_of_joining']} to {data['date_of_completion']}",
        "generated_at":      datetime.now().strftime("%d %b %Y, %I:%M %p"),
        "departments_count": 1,
        "master_used":       "Appraisal Proforma Master"
    })

    return send_file(filepath, as_attachment=True)


# COORDINATOR NOMINATION GENERATOR
# ══════════════════════════════════════
@app.route("/coordinator-nomination/<master_id>", methods=["GET"])
def coordinator_nomination(master_id):
    defaults = load_defaults()
    masters = load_masters()
    master = next((m for m in masters if m["id"] == master_id), None)
    if not master:
        flash(f"Master template {master_id} not found.")
        return redirect(url_for("internship_menu"))
        
    groups = ["AFTD", "ADS", "BIDS", "BEHI", "SS", "TELIC", "PPG", "QMG", "ETF", "PC", "WHD", "EXPD", "WHT&E", "ARISE", "PCD", "AIG", "RTRS", "S&D", "R&QA", "WKS", "SEED", "CERBERUS", "DPB", "HSP", "I2G", "HRDD", "BTS"]
    return render_template("coordinator_nomination_form.html", defaults=defaults, groups=groups, master=master)


@app.route("/generate-coordinator-nomination/<master_id>", methods=["POST"])
def generate_coordinator_nomination_route(master_id):
    masters = load_masters()
    master = next((m for m in masters if m["id"] == master_id), None)
    if not master:
        return "Master not found", 404
        
    trainees_json = form_get("trainees_json", "[]")
    try:
        trainees = json.loads(trainees_json)
    except Exception:
        trainees = []

    data = {
        "training_session": form_get("training_session", ""),
        "ref_no":           form_get("ref_no", ""),
        "date":             form_get("date", ""),
        "signatory_name":   form_get("signatory_name", ""),
        "signatory_desig":  form_get("signatory_desig", ""),
        "recipient_group":  form_get("recipient_group", ""),
        "trainees":         trainees
    }

    filepath, filename = generate_coordinator_nomination(master_id, data)

    save_history({
        "filename":          filename,
        "degree":            f"Selected Trainees ({len(trainees)})",
        "period":            data["training_session"],
        "generated_at":      datetime.now().strftime("%d %b %Y, %I:%M %p"),
        "departments_count": len(trainees),
        "master_used":       master["name"]
    })

    return send_file(filepath, as_attachment=True)


# INTERNSHIP NOTING GENERATOR (NEW MULTI-TABLE ROUTE)
# ══════════════════════════════════════
@app.route("/internship-noting/<master_id>", methods=["GET"])
def internship_noting(master_id):
    defaults = load_defaults()
    masters = load_masters()
    master = next((m for m in masters if m["id"] == master_id), None)
    if not master:
        flash(f"Master template {master_id} not found.")
        return redirect(url_for("internship_menu"))
        
    groups = ["AFTD", "ADS", "BIDS", "BEHI", "SS", "TELIC", "PPG", "QMG", "ETF", "PC", "WHD", "EXPD", "WHT&E", "ARISE", "PCD", "AIG", "RTRS", "S&D", "R&QA", "WKS", "SEED", "CERBERUS", "DPB", "HSP", "I2G", "HRDD", "BTS"]
    return render_template("internship_grid_form.html", defaults=defaults, groups=groups, master=master)


@app.route("/generate-internship-noting/<master_id>", methods=["POST"])
def generate_internship_noting_route(master_id):
    masters = load_masters()
    master = next((m for m in masters if m["id"] == master_id), None)
    if not master:
        return "Master not found", 404
        
    rows_json = form_get("rows_json", "[]")
    try:
        rows = json.loads(rows_json)
    except Exception:
        rows = []

    data = {"rows": rows}
    for var in master.get("variables", []):
        data[var.lower()] = form_get(var, "")

    filepath, filename = generate_internship_noting(master_id, data)

    save_history({
        "filename":          filename,
        "degree":            f"Dynamic Grid ({len(rows)} rows)",
        "period":            data.get("date", "N/A"),
        "generated_at":      datetime.now().strftime("%d %b %Y, %I:%M %p"),
        "departments_count": len(rows),
        "master_used":       master["name"]
    })

    return send_file(filepath, as_attachment=True)


# ── API History ──
@app.route("/api/history")
def api_history():
    from flask import jsonify
    return jsonify({"history": load_history()})

# ── Restructured Module Routes (Training & Internship) ──
@app.route("/training")
def training_menu():
    """Renders the Training Module submenu with pillars."""
    category = request.args.get("sub_module", "noting")
    return render_template("training_menu.html", active_category=category)

@app.route("/internship")
def internship_menu():
    """Renders the Internship Module submenu."""
    return render_template("internship_menu.html")

@app.route("/api/templates")
def api_templates():
    """Returns templates filtered by module and sub_module."""
    module = request.args.get("module")
    sub_module = request.args.get("sub_module")
    masters = load_masters()
    
    filtered = masters
    if module:
        filtered = [m for m in filtered if m.get("module") == module]
    if sub_module:
        filtered = [m for m in filtered if m.get("sub_module") == sub_module]
        
    return jsonify(filtered)

if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5001, debug=False)